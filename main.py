#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
网络安全值守保障日报整合脚本
功能：1. IP归属分析 2. 值守日报生成
依赖：pip install pandas openpyxl python-docx requests py-ip2region
"""

import os
import re
import sys
import ipaddress
import subprocess
import time
import warnings
import threading
from pathlib import Path
from datetime import datetime
from tkinter import (
    Tk, Frame, Label, Button, Listbox, Entry, Text,
    Scrollbar, filedialog, messagebox, StringVar, END, NONE,
    DISABLED, NORMAL
)

# 修复 numpy 2.x 在打包后 DLL 加载问题
def _setup_dll_paths():
    """在打包模式下设置 numpy/pandas 的 DLL 路径"""
    if not getattr(sys, 'frozen', False):
        return
    
    # 获取基础目录
    if hasattr(sys, '_MEIPASS'):
        base_dir = sys._MEIPASS
    else:
        base_dir = os.path.dirname(sys.executable)
    
    # 添加 numpy.libs 目录
    numpy_libs = os.path.join(base_dir, 'numpy.libs')
    if os.path.isdir(numpy_libs):
        try:
            os.add_dll_directory(numpy_libs)
        except (OSError, AttributeError):
            if hasattr(os, 'environ'):
                os.environ['PATH'] = numpy_libs + os.pathsep + os.environ.get('PATH', '')
    
    # 添加 pandas.libs 目录
    pandas_libs = os.path.join(base_dir, 'pandas.libs')
    if os.path.isdir(pandas_libs):
        try:
            os.add_dll_directory(pandas_libs)
        except (OSError, AttributeError):
            if hasattr(os, 'environ'):
                os.environ['PATH'] = pandas_libs + os.pathsep + os.environ.get('PATH', '')
    
    # 调试信息（可选）
    if os.environ.get('DEBUG_NUMPY', '0') == '1':
        print(f"[DEBUG] frozen={getattr(sys, 'frozen', False)}")
        print(f"[DEBUG] base_dir={base_dir}")
        print(f"[DEBUG] numpy.libs exists={os.path.isdir(numpy_libs)}")
        print(f"[DEBUG] pandas.libs exists={os.path.isdir(pandas_libs)}")
        if os.path.isdir(numpy_libs):
            print(f"[DEBUG] numpy.libs contents: {os.listdir(numpy_libs)}")
        if os.path.isdir(pandas_libs):
            print(f"[DEBUG] pandas.libs contents: {os.listdir(pandas_libs)}")

_setup_dll_paths()

warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')

REQUIRED = {'pandas': 'pandas', 'openpyxl': 'openpyxl', 'python-docx': 'docx', 'requests': 'requests', 'py-ip2region': 'ip2region'}

def ensure_deps():
    # 打包后禁用自动安装（sys.executable会指向exe本身，导致无限递归）
    # 同时检查 frozen 和 _MEIPASS，确保在各种打包场景下都能正确检测
    if getattr(sys, 'frozen', False) or hasattr(sys, '_MEIPASS'):
        return
    if os.environ.get('NO_AUTO_INSTALL') == '1':
        return
    miss = [p for p, i in REQUIRED.items() if _import_failed(i)]
    if not miss:
        return
    print(f'[!] 自动安装缺失依赖: {miss}')
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', *miss,
                           '-i', 'https://pypi.tuna.tsinghua.edu.cn/simple'])

def _import_failed(name):
    try:
        __import__(name)
        return False
    except ImportError:
        return True

ensure_deps()

# ================ 版本 & 自动更新 ================
APP_VERSION = "1.3.0"

# 版本信息来源：按优先级依次尝试
# 在 GitHub / 内网HTTP / 共享目录 放置 version.json，例如：
# {
#   "version": "1.1.0",
#   "exe_urls": [
#     "https://github.com/eavil666/dayupdate/releases/download/v1.1.0/daily-report.exe",
#     "https://ghfast.top/https://github.com/eavil666/dayupdate/releases/download/v1.1.0/daily-report.exe",
#     "https://ghproxy.net/https://github.com/eavil666/dayupdate/releases/download/v1.1.0/daily-report.exe"
#   ],
#   "md5": "abc123...",
#   "release_note": "修复IP归属显示问题",
#   "force_update": false
# }
# exe_urls 为数组时按顺序回退；也兼容 exe_url 单字符串格式
# GitHub API：始终返回真实最新 release（无 CDN 缓存），作为首选版本源
# 注意：未鉴权时 60 次/小时限频，桌面应用每次启动查一次足够
GITHUB_API_LATEST = "https://api.github.com/repos/eavil666/dayupdate/releases/latest"

# EXE 下载 CDN 镜像前缀（应用到 GitHub 直链前，按顺序回退）
CDN_MIRROR_PREFIXES = [
    "",  # GitHub 直链
    "https://ghfast.top/",
    "https://ghproxy.net/",
    "https://gh-proxy.com/",
    "https://mirror.ghproxy.com/",
]

# version.json 兜底源（GitHub API 失败时使用）
# 注意：不要用 jsDelivr（cdn.jsdelivr.net）——它对 version.json 长缓存且忽略 ?t= 参数，
# 会返回远古旧版本（实测缓存到 1.0.0），导致"一版一版升"或检测错乱
UPDATE_VERSION_URLS = [
    # raw 优先：缓存仅 5 分钟
    "https://raw.githubusercontent.com/eavil666/dayupdate/main/version.json",
    # ghproxy/ghfast 镜像 raw：国内可达，无长缓存问题
    "https://ghproxy.net/https://raw.githubusercontent.com/eavil666/dayupdate/main/version.json",
    "https://ghfast.top/https://raw.githubusercontent.com/eavil666/dayupdate/main/version.json",
    # "http://intranet-server/apps/report/version.json",
    # r"\\file-server\share\report\version.json",
]

# 备选EXE下载源（如果version.json内未提供exe_urls/exe_url）
UPDATE_EXE_URLS = [
    "https://github.com/eavil666/dayupdate/releases/download/v{version}/daily-report.exe",
    "https://ghfast.top/https://github.com/eavil666/dayupdate/releases/download/v{version}/daily-report.exe",
    # "http://intranet-server/apps/report/daily-report-{version}.exe",
    # r"\\file-server\share\report\daily-report-{version}.exe",
]


def _parse_version(v):
    parts = []
    for x in re.findall(r'\d+', str(v)):
        try:
            parts.append(int(x))
        except (ValueError, TypeError):
            parts.append(0)
    return tuple(parts)


class AutoUpdater:
    def __init__(self, current_version=APP_VERSION,
                 version_urls=None, exe_urls=None,
                 progress_cb=None, log_cb=None, ask_confirm_cb=None):
        self.current_version = current_version
        self.version_urls = version_urls or UPDATE_VERSION_URLS
        self.exe_urls = exe_urls or UPDATE_EXE_URLS
        self.progress_cb = progress_cb or (lambda v, m: None)
        self.log_cb = log_cb or (lambda m: print(m))
        self.ask_confirm_cb = ask_confirm_cb or (lambda msg: True)
        self._latest_info = None

    def _fetch_latest_release_api(self):
        """通过 GitHub API 获取最新 release（无 CDN 缓存，始终返回真实最新版）。

        解决 jsDelivr 等对 version.json 长缓存导致的"一版一版升"问题：
        国内 raw.githubusercontent.com 常不通，程序会降级到 jsDelivr，而 jsDelivr
        缓存可能停留在远古版本（实测缓存到 1.0.0），导致检测到的"最新版"是旧版。
        GitHub API 无此问题——每次请求都返回真实最新 release。

        返回与 version.json 同结构的 dict；API 不提供 md5，下载时不做 MD5 校验
        （requests 流式下载 + raise_for_status 已能捕获传输错误，足够可靠）。
        """
        import requests
        try:
            r = requests.get(GITHUB_API_LATEST, timeout=10,
                             headers={'User-Agent': 'daily-report-updater'})
            r.raise_for_status()
            data = r.json()
            tag = data.get('tag_name', '') or ''
            version = tag.lstrip('vV').strip()
            # 找到 .exe 资产
            asset_url = None
            for a in data.get('assets', []) or []:
                if (a.get('name') or '').lower().endswith('.exe'):
                    asset_url = a.get('browser_download_url')
                    break
            if not version or not asset_url:
                return None
            # 用 CDN 镜像前缀构造回退下载列表（直链在前，镜像兜底）
            exe_urls = [p + asset_url for p in CDN_MIRROR_PREFIXES]
            self.log_cb(f"[版本] GitHub API 获取最新版本成功 (version={version}, tag={tag})")
            return {
                'version': version,
                'exe_urls': exe_urls,
                'md5': None,  # API 不提供 md5，下载时跳过 MD5 校验
                'release_note': data.get('body', '') or '',
                'force_update': False,
            }
        except Exception as exc:
            self.log_cb(f"[版本] GitHub API 获取失败: {exc}")
            return None

    def _fetch_version_json(self):
        if not self.version_urls:
            return None
        for url in self.version_urls:
            try:
                if url.startswith('http://') or url.startswith('https://'):
                    import requests
                    # CDN 防缓存：加 ?t=时间戳（raw 本身不缓存；ghproxy 镜像也不长缓存）
                    if '?' in url:
                        fetch_url = f"{url}&t={int(time.time())}"
                    else:
                        fetch_url = f"{url}?t={int(time.time())}"
                    r = requests.get(fetch_url, timeout=10)
                    r.raise_for_status()
                    data = r.json()
                else:
                    # 共享目录 / 本地文件
                    import json
                    with open(url, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                if isinstance(data, dict) and 'version' in data:
                    self.log_cb(f"[版本] 从 {url} 获取版本信息成功 (version={data.get('version')})")
                    return data
            except Exception as exc:
                self.log_cb(f"[版本] {url} 获取失败: {exc}")
        return None

    def check_update(self):
        # 首选 GitHub API：无 CDN 缓存，始终返回真实最新版，
        # 避免 jsDelivr/CDN 长缓存导致的"一版一版升"或检测错乱
        data = self._fetch_latest_release_api()
        if not data:
            # 降级到 version.json 源（raw + ghproxy 镜像，已剔除 jsDelivr）
            data = self._fetch_version_json()
        if not data:
            self.log_cb("[版本] 未配置更新源或获取失败，跳过")
            return None
        latest_ver = data.get('version', '0.0.0')
        self.log_cb(f"[版本] 当前版本 v{self.current_version} | 最新版本 v{latest_ver}")
        if _parse_version(latest_ver) > _parse_version(self.current_version):
            self.log_cb(f"[版本] 发现新版本 v{latest_ver}，准备更新")
            self._latest_info = data
            return data
        self.log_cb(f"[版本] 当前版本 v{self.current_version} 已是最新")
        return None

    def _download(self, url, dest_path, expected_md5=None):
        file_size = 0
        if url.startswith('http://') or url.startswith('https://'):
            import requests
            resp = requests.get(url, timeout=(30, 600), stream=True)
            resp.raise_for_status()
            total = int(resp.headers.get('content-length', 0))
            if total > 0:
                self.progress_cb(0, total)
            downloaded = 0
            with open(dest_path, 'wb') as f:
                for chunk in resp.iter_content(chunk_size=1024 * 1024):
                    if chunk:
                        f.write(chunk)
                        downloaded += len(chunk)
                        if total > 0:
                            self.progress_cb(downloaded, total)
            file_size = downloaded
        else:
            # 共享目录 / 本地文件
            total = os.path.getsize(url)
            if total > 0:
                self.progress_cb(0, total)
            import shutil
            shutil.copy2(url, dest_path)
            self.progress_cb(total, total)
            file_size = total

        # MD5 校验
        if expected_md5:
            import hashlib
            h = hashlib.md5()
            with open(dest_path, 'rb') as f:
                for chunk in iter(lambda: f.read(1024 * 1024), b''):
                    h.update(chunk)
            actual = h.hexdigest().lower()
            if actual != expected_md5.lower():
                raise RuntimeError(f"MD5校验失败: 期望 {expected_md5}, 实际 {actual}")
            self.log_cb(f"[更新] MD5校验通过: {actual}")
        self.progress_cb(0, 0)
        return file_size

    def _resolve_exe_urls(self, info):
        """解析 EXE 下载地址列表，支持 exe_urls(数组/CDN回退) 和 exe_url(单字符串/兼容)"""
        if info:
            if info.get('exe_urls'):
                return info['exe_urls']
            if info.get('exe_url'):
                return [info['exe_url']]
        version = info.get('version', '') if info else ''
        urls = []
        for url in self.exe_urls:
            try:
                urls.append(url.format(version=version))
            except (KeyError, IndexError):
                continue
        return urls

    def download_update(self, info=None):
        if info is None:
            info = self._latest_info or self.check_update()
        if not info:
            return None

        exe_urls = self._resolve_exe_urls(info)
        if not exe_urls:
            self.log_cb("[更新] 未配置EXE下载地址")
            return None

        import tempfile
        tmp_dir = tempfile.gettempdir()
        tmp_exe = os.path.join(tmp_dir, f"report_update_{int(time.time())}.exe")

        for idx, exe_url in enumerate(exe_urls):
            try:
                self.log_cb(f"[更新] 下载新版本 v{info.get('version')}: {exe_url}")
                size = self._download(exe_url, tmp_exe, info.get('md5'))
                self.log_cb(f"[更新] 下载完成: {size // 1024 // 1024} MB")
                return tmp_exe
            except Exception as exc:
                self.log_cb(f"[更新] 下载失败: {exc}")
                try:
                    if os.path.exists(tmp_exe):
                        os.remove(tmp_exe)
                except OSError as e:
                    self.log_cb(f'[!] 清理临时文件失败: {e}')
                if idx < len(exe_urls) - 1:
                    self.log_cb("[更新] 尝试下一个镜像源...")
                else:
                    self.log_cb("[!] 所有镜像源均下载失败")
                    return None

    def install_and_restart(self, new_exe_path):
        if not new_exe_path or not os.path.exists(new_exe_path):
            return False
        if not getattr(sys, 'frozen', False):
            self.log_cb("[更新] 源码模式跳过安装（仅EXE模式支持自动替换）")
            return False

        old_exe = os.path.abspath(sys.executable)
        old_dir = os.path.dirname(old_exe)
        backup_exe = old_exe + ".bak"
        parent_pid = os.getpid()

        # === 用当前 exe 自身作为 Worker 模式更新 ===
        # 历史踩坑：
        #   - .bat：中文路径在 GBK/UTF-8 codepage 间乱码 → 不替换/生成乱码文件
        #   - PowerShell：5.1 对 PyInstaller frozen 父进程触发安全校验弹窗
        #   - 外部 updater.exe：循环依赖（旧版本不含 updater.exe 无法更新到新版本）
        # 结论：用自身 --update-worker 模式，任何版本都能自更新，无外部依赖。
        # worker 模式入口在 main() 顶部，检测到 --update-worker=<jsonPath> 即
        # 跳到 update_worker_main()，不加载任何 GUI/Tkinter。
        import json, tempfile, hashlib

        tag = hashlib.md5(str(time.time()).encode()).hexdigest()[:8]
        json_params_path = os.path.join(tempfile.gettempdir(), f"update_{tag}_params.json")

        params_json = json.dumps({
            'parentPid': parent_pid,
            'oldExe':  old_exe,
            'newExe':  new_exe_path,
            'bakExe':  backup_exe,
            'workDir': old_dir,
            'jsonPath': json_params_path,
        }, ensure_ascii=False, indent=2)

        try:
            with open(json_params_path, 'w', encoding='utf-8') as f:
                f.write(params_json)
        except Exception as exc:
            self.log_cb(f"[更新] 写入参数文件失败: {exc}")
            return False

        # 启动 Worker（DETACHED，独立于父进程；os._exit 后 worker 继续运行）
        CREATE_NO_WINDOW = 0x08000000
        DETACHED_PROCESS = 0x00000008
        CREATE_NEW_PROCESS_GROUP = 0x00000200
        flags = CREATE_NO_WINDOW | DETACHED_PROCESS | CREATE_NEW_PROCESS_GROUP

        worker_arg = f'--update-worker={json_params_path}'
        try:
            subprocess.Popen(
                [sys.executable, worker_arg],
                cwd=old_dir,
                creationflags=flags,
            )
            self.log_cb(f"[更新] 启动更新 Worker (pid={parent_pid} -> worker)")
        except OSError as e:
            self.log_cb(f"[!] 启动更新 Worker 失败: {e}")
            return False

        self.log_cb(f"[更新] 日志将写入: %TEMP%\\update_last.log 和 {os.path.join(old_dir, 'update_last.log')}")
        self.log_cb("[更新] 正在重启应用完成更新...")
        os._exit(0)

    def run_update_flow(self, force_dialog=False):
        info = self.check_update()
        if not info:
            if force_dialog:
                self.log_cb("[更新] 当前已是最新版本")
            return False
        ver = info.get('version', '?')
        note = info.get('release_note', '')
        force = bool(info.get('force_update', False))

        prompt = f"发现新版本 v{ver}\n\n当前版本: v{self.current_version}\n新版本: v{ver}"
        if note:
            prompt += f"\n\n更新说明:\n{note}"
        prompt += "\n\n是否立即更新？"

        if force or self.ask_confirm_cb(prompt):
            tmp = self.download_update(info)
            if tmp:
                return self.install_and_restart(tmp)
        return False


import configparser

script_dir = os.path.dirname(os.path.abspath(__file__))
# 运行时基础目录（打包后优先使用exe所在目录）
if getattr(sys, 'frozen', False):
    # 单文件模式：优先使用exe所在目录，用户的数据文件放在这里
    exe_dir = os.path.dirname(sys.executable)
    runtime_dir = exe_dir
    # 记录临时解压目录，用于读取内置资源
    if hasattr(sys, '_MEIPASS'):
        meipass_dir = sys._MEIPASS
    else:
        meipass_dir = None
else:
    runtime_dir = script_dir
    meipass_dir = None

def _load_excluded_ip_networks(config_path: str = None) -> list:
    if config_path is None:
        config_path = os.path.join(runtime_dir, 'config.ini')
    config = configparser.ConfigParser()
    config.read(config_path, encoding='utf-8')
    excluded_ips_str = config.get('network', 'excluded_ips', fallback='')
    networks = []
    for _ex_range in (r.strip() for r in excluded_ips_str.split(',') if r.strip()):
        if '-' not in _ex_range:
            try:
                networks.append(ipaddress.ip_address(_ex_range))
            except ValueError:
                pass
            continue
        parts = _ex_range.split('-')
        if len(parts) != 2:
            continue
        base, end_part = parts[0].strip(), parts[1].strip()
        # 简写展开：1.2.3.4-10 → 1.2.3.4-1.2.3.10（与 _parse_ip_range 行为一致）
        end_parts = end_part.split('.')
        if len(end_parts) == 1:
            base_parts = base.split('.')
            if len(base_parts) == 4:
                end_part = f'{base_parts[0]}.{base_parts[1]}.{base_parts[2]}.{end_parts[0]}'
            else:
                continue
        try:
            start_ip = ipaddress.ip_address(base)
            end_ip = ipaddress.ip_address(end_part)
            if int(start_ip) > int(end_ip):
                start_ip, end_ip = end_ip, start_ip
            networks.extend(ipaddress.summarize_address_range(start_ip, end_ip))
        except ValueError:
            pass
    return networks

EXCLUDED_IP_NETWORKS = _load_excluded_ip_networks()
EXCLUDED_IP_LABELS = {}  # IP -> 说明（来自外部Excel）

def is_excluded_ip(ip_str):
    try:
        ip = ipaddress.ip_address(str(ip_str).strip())
    except ValueError:
        return False
    for net in EXCLUDED_IP_NETWORKS:
        if isinstance(net, ipaddress.IPv4Address):
            if ip == net:
                return True
        else:
            if ip in net:
                return True
    return False

def load_external_excluded_ips(excel_path):
    """从外部Excel文件加载排除业务IP，格式：ip, 说明"""
    global EXCLUDED_IP_NETWORKS, EXCLUDED_IP_LABELS
    import pandas as pd
    try:
        df = pd.read_excel(excel_path)
        df.columns = df.columns.str.strip()
        # 识别IP列和说明列
        ip_col = None
        label_col = None
        for col in df.columns:
            col_lower = str(col).lower()
            if 'ip' in col_lower and ip_col is None:
                ip_col = col
            elif '说明' in str(col) or '备注' in str(col) or 'desc' in col_lower:
                label_col = col
        if ip_col is None:
            # 如果没找到带ip的列名，取第一列作为IP
            ip_col = df.columns[0]
            if label_col is None and len(df.columns) > 1:
                label_col = df.columns[1]
        count = 0
        for _, row in df.iterrows():
            ip_val = row[ip_col]
            if pd.isna(ip_val):
                continue
            ip_str = str(ip_val).strip()
            label_str = str(row[label_col]).strip() if label_col and not pd.isna(row[label_col]) else '业务IP'
            # 复用 _parse_ip_range 统一处理单IP/范围/简写
            ips = _parse_ip_range(ip_str)
            if not ips:
                continue
            start_ip = ipaddress.ip_address(ips[0])
            end_ip = ipaddress.ip_address(ips[-1])
            if len(ips) == 1:
                EXCLUDED_IP_NETWORKS.append(start_ip)
            else:
                for net in ipaddress.summarize_address_range(start_ip, end_ip):
                    EXCLUDED_IP_NETWORKS.append(net)
            for ip in ips:
                EXCLUDED_IP_LABELS[ip] = label_str
            count += 1
        _log(f'[+] 从外部文件加载业务IP: {count} 条')
        return count
    except Exception as e:
        _log(f'[!] 加载外部业务IP文件失败: {e}')
        return 0

def is_private_ip(ip_str):
    try:
        ip = ipaddress.ip_address(str(ip_str).strip())
        return ip.is_private or ip.is_loopback or ip.is_reserved or ip.is_link_local or ip.is_multicast
    except ValueError:
        return False

TERMINAL_IP_TABLE = None
TERMINAL_IP_TABLE_PATH = None  # 外部导入的终端IP表路径

def set_terminal_ip_table_path(path):
    """设置终端IP地址表路径并重置缓存，供GUI导入按钮调用"""
    global TERMINAL_IP_TABLE, TERMINAL_IP_TABLE_PATH
    TERMINAL_IP_TABLE_PATH = path
    TERMINAL_IP_TABLE = None  # 重置缓存，下次调用时重新加载

def _parse_ip_range(ip_str):
    """解析IP字符串，支持单IP和范围格式（完整/简写），返回IP列表。
    支持格式:
      - 单IP: '172.16.70.226'
      - 完整范围: '172.16.70.226-172.16.70.230'
      - 简写范围: '172.16.70.226-230'
    异常时返回空列表并打印警告日志。
    """
    result = []
    ip_str = str(ip_str).strip()
    if not ip_str:
        return result
    try:
        if '-' not in ip_str:
            # 单IP
            ipaddress.ip_address(ip_str)  # 校验格式
            result.append(ip_str)
            return result

        parts = ip_str.split('-')
        if len(parts) != 2:
            _log(f'[!] 跳过无效IP格式（包含多个"-"）: {ip_str}')
            return result

        start_str = parts[0].strip()
        end_str = parts[1].strip()
        if not start_str or not end_str:
            _log(f'[!] 跳过无效IP范围（起止IP为空）: {ip_str}')
            return result

        start_ip = ipaddress.ip_address(start_str)

        # 处理简写格式如 172.16.70.226-230
        end_parts = end_str.split('.')
        if len(end_parts) == 1:
            base_parts = start_str.split('.')
            if len(base_parts) == 4:
                end_str = f'{base_parts[0]}.{base_parts[1]}.{base_parts[2]}.{end_parts[0]}'
            else:
                _log(f'[!] 跳过无效IP范围（简写格式无法推断前三段）: {ip_str}')
                return result

        end_ip = ipaddress.ip_address(end_str)

        start_int = int(start_ip)
        end_int = int(end_ip)
        if start_int > end_int:
            # 自动交换，兼容 end-start 书写
            start_int, end_int = end_int, start_int

        span = end_int - start_int + 1
        # 大范围保护：超过 65536 (B类整段) 截断并警告，避免 OOM
        MAX_IP_RANGE = 65536
        if span > MAX_IP_RANGE:
            _log(f'[!] IP范围过大 ({span} 个地址)，截断为前 {MAX_IP_RANGE} 个: {ip_str}')
            end_int = start_int + MAX_IP_RANGE - 1
            span = MAX_IP_RANGE

        for ip_int in range(start_int, end_int + 1):
            result.append(str(ipaddress.ip_address(ip_int)))

    except ValueError as e:
        _log(f'[!] 跳过无效IP: {ip_str} ({e})')
    except Exception as e:
        _log(f'[!] 解析IP异常: {ip_str} ({e})')

    return result

def load_terminal_ip_table():
    global TERMINAL_IP_TABLE
    if TERMINAL_IP_TABLE is not None:
        return TERMINAL_IP_TABLE
    TERMINAL_IP_TABLE = {}
    # 优先使用外部导入路径，其次从文件系统查找
    ip_table_path = TERMINAL_IP_TABLE_PATH or _find_file("终端ip地址表.xlsx")
    if not ip_table_path or not os.path.exists(ip_table_path):
        return TERMINAL_IP_TABLE
    try:
        import pandas as pd
        df = pd.read_excel(ip_table_path)
        skipped = 0
        for _, row in df.iterrows():
            ip_value = row.iloc[2]
            dept_value = row.iloc[1] if len(row) > 1 else ''
            if pd.isna(ip_value):
                continue
            ip_str = str(ip_value).strip()
            dept_str = str(dept_value).strip() if not pd.isna(dept_value) else ''
            ips = _parse_ip_range(ip_str)
            if not ips:
                skipped += 1
                continue
            for ip in ips:
                TERMINAL_IP_TABLE[ip] = dept_str
        if skipped:
            _log(f'[*] 终端IP表: 跳过 {skipped} 行无效记录')
    except Exception as e:
        _log(f'[!] 加载终端IP表异常: {e}')
    return TERMINAL_IP_TABLE

def get_terminal_location(ip):
    ip = str(ip).strip()
    table = load_terminal_ip_table()
    return table.get(ip, '')

def is_valid_public_ip(ip):
    try:
        addr = ipaddress.ip_address(ip)
        return not (addr.is_private or addr.is_loopback or addr.is_reserved or addr.is_link_local or addr.is_multicast)
    except ValueError:
        return False

def local_ip_label(ip):
    try:
        ipaddress.ip_address(ip)
    except ValueError:
        return "无效IP"
    if not is_valid_public_ip(ip):
        return "内网/保留地址"
    return None

def format_online_result(data):
    if data.get("status") != "success":
        return data.get("message") or "查询失败"
    parts = []
    for key in ("country", "regionName", "city", "isp"):
        val = (data.get(key) or "").strip()
        if val:
            parts.append(val)
    return " ".join(parts) if parts else "未知"

BATCH_SIZE = 100
BATCH_INTERVAL = 1.5
XDB_FILE = None
DOWNLOAD_SOURCES = [
    "https://edgeone.gh-proxy.com/https://github.com/lionsoul2014/ip2region/raw/refs/heads/master/data/ip2region_v4.xdb",
    "https://gcore.jsdelivr.net/gh/lionsoul2014/ip2region/data/ip2region_v4.xdb",
    "https://raw.githubusercontent.com/lionsoul2014/ip2region/master/data/ip2region_v4.xdb",
    "https://gitee.com/lionsoul/ip2region/raw/master/data/ip2region_v4.xdb",
]

# GUI 日志回调，由 App 实例注册
_gui_log_callback = None
_gui_progress_callback = None

def _log(msg):
    """输出日志：GUI 模式发到界面，否则 print"""
    if _gui_log_callback:
        _gui_log_callback(msg)
    else:
        print(msg)

def _set_progress(value, maximum=None):
    """设置进度条：value 为当前值，maximum 为最大值（None 表示不确定模式）"""
    if _gui_progress_callback:
        _gui_progress_callback(value, maximum)

def download_xdb():
    import requests
    for idx, url in enumerate(DOWNLOAD_SOURCES, 1):
        try:
            _log(f"[{idx}/{len(DOWNLOAD_SOURCES)}] 正在下载离线 IP 库: {url}")
            resp = requests.get(url, timeout=(30, 600), stream=True)
            resp.raise_for_status()
            total = int(resp.headers.get('content-length', 0))
            if total > 0:
                _set_progress(0, total)
            downloaded = 0
            with open(XDB_FILE, "wb") as f:
                for chunk in resp.iter_content(chunk_size=1024 * 1024):
                    if chunk:
                        f.write(chunk)
                        downloaded += len(chunk)
                        if total > 0:
                            _set_progress(downloaded, total)
            size_kb = os.path.getsize(XDB_FILE) // 1024
            if size_kb < 9000:
                _log(f"文件大小异常({size_kb}KB)，尝试下一个地址")
                os.remove(XDB_FILE)
                continue
            _log(f"已保存: {XDB_FILE} ({size_kb} KB)")
            _set_progress(0, 0)
            return True
        except Exception as exc:
            _log(f"下载失败: {exc}")
            _set_progress(0, 0)
    _log("所有下载地址均失败，将使用在线 API 查询")
    _set_progress(0, 0)
    return False

def _get_requests_verify():
    """返回 requests 可用的 verify 参数：frozen 模式下用 certifi 自带 CA，失败回退 False"""
    try:
        if getattr(sys, 'frozen', False):
            try:
                import certifi
                ca_path = certifi.where()
                if os.path.exists(ca_path):
                    # 让 requests 使用 certifi 的 CA bundle（PyInstaller frozen 模式找不到系统证书）
                    os.environ.setdefault('REQUESTS_CA_BUNDLE', ca_path)
                    os.environ.setdefault('SSL_CERT_FILE', ca_path)
                    return ca_path
            except Exception:
                pass
        return True
    except Exception:
        return True

def download_xdb():
    import requests
    verify = _get_requests_verify()
    for idx, url in enumerate(DOWNLOAD_SOURCES, 1):
        try:
            _log(f"[{idx}/{len(DOWNLOAD_SOURCES)}] 正在下载离线 IP 库: {url}")
            # 第一次尝试使用 verify（正常CA验证）
            try:
                resp = requests.get(url, timeout=(30, 600), stream=True, verify=verify)
                resp.raise_for_status()
            except Exception as exc:
                # 如果是 SSL 相关错误且 verify 不是 False，则降级 verify=False 重试
                is_ssl_err = ('SSL' in str(type(exc).__name__) or 'SSL' in str(exc)
                              or 'CERTIFICATE' in str(exc).upper()
                              or isinstance(exc, requests.exceptions.SSLError))
                if is_ssl_err and verify is not False:
                    _log(f"  SSL验证失败，跳过证书验证重试...")
                    resp = requests.get(url, timeout=(30, 600), stream=True, verify=False)
                    resp.raise_for_status()
                    # 首次 SSL 降级成功后，后续镜像也统一用 verify=False 提速
                    verify = False
                else:
                    raise
            total = int(resp.headers.get('content-length', 0))
            if total > 0:
                _set_progress(0, total)
            downloaded = 0
            # 确保目标目录存在（exe所在目录可能需要创建）
            os.makedirs(os.path.dirname(os.path.abspath(XDB_FILE)), exist_ok=True)
            with open(XDB_FILE, "wb") as f:
                for chunk in resp.iter_content(chunk_size=1024 * 1024):
                    if chunk:
                        f.write(chunk)
                        downloaded += len(chunk)
                        if total > 0:
                            _set_progress(downloaded, total)
            size_kb = os.path.getsize(XDB_FILE) // 1024
            if size_kb < 9000:
                _log(f"文件大小异常({size_kb}KB)，尝试下一个地址")
                try:
                    os.remove(XDB_FILE)
                except OSError:
                    pass
                continue
            _log(f"已保存: {XDB_FILE} ({size_kb} KB)")
            _set_progress(0, 0)
            return True
        except Exception as exc:
            _log(f"下载失败: {exc}")
            _set_progress(0, 0)
    _log("所有下载地址均失败，将使用在线 API 查询")
    _set_progress(0, 0)
    return False

def get_offline_searcher():
    global XDB_FILE
    if XDB_FILE is None:
        # 优先使用exe所在目录（打包后）
        if getattr(sys, 'frozen', False):
            XDB_FILE = os.path.join(os.path.dirname(sys.executable), "ip2region_v4.xdb")
        else:
            XDB_FILE = os.path.join(script_dir, "ip2region_v4.xdb")
    _log(f"[?] 检查本地IP库文件: {XDB_FILE}")
    # 检查文件是否存在及是否超过7天需要更新
    need_download = False
    if not os.path.exists(XDB_FILE):
        _log("[!] 本地IP库不存在，开始下载...")
        need_download = True
    else:
        file_mtime = os.path.getmtime(XDB_FILE)
        file_age_days = (time.time() - file_mtime) / 86400
        if file_age_days >= 7:
            _log(f"[!] 离线IP库已超过{int(file_age_days)}天，自动更新...")
            need_download = True
    if need_download:
        if not download_xdb():
            return None
    if not os.path.exists(XDB_FILE):
        _log("[!] 离线IP库文件不存在，将使用在线API查询")
        return None
    try:
        import ip2region.searcher as xdb
        import ip2region.util as util
        searcher = xdb.new_with_file_only(util.IPv4, str(XDB_FILE))
        _log("[+] 离线IP库加载成功")
        return searcher
    except Exception as exc:
        _log(f"[!] 离线IP库加载失败: {exc}，将使用在线API查询")
        return None

def parse_region(raw):
    parts = (raw or "").split("|")
    while len(parts) < 5:
        parts.append("")
    country, province, city, isp, _code = parts[:5]
    display = []
    if country and country not in ("0", "内网IP"):
        if country != "中国":
            display.append(country)
    if province and province not in ("0", ""):
        display.append(province)
    if city and city not in ("0", ""):
        display.append(city)
    if isp and isp not in ("0", ""):
        display.append(isp)
    location = " ".join(display) if display else (raw or "未知")
    stat_province = province if province not in ("0", "") else country
    stat_city = city if city not in ("0", "") else stat_province
    if not stat_province:
        stat_province = "未知"
    if not stat_city:
        stat_city = stat_province
    return location, stat_province, stat_city

def query_offline(searcher, ips):
    results = {}
    for ip in ips:
        if not is_valid_public_ip(ip):
            continue
        try:
            raw = searcher.search(ip)
            if not raw:
                results[ip] = ("未知", "未知", "未知")
            else:
                results[ip] = parse_region(raw)
        except Exception as exc:
            results[ip] = (f"查询失败({exc})", "查询失败", "查询失败")
    return results

def query_online_batch(ips):
    import requests
    results = {}
    public_ips = [ip for ip in ips if is_valid_public_ip(ip)]
    if not public_ips:
        return results
    url = "http://ip-api.com/batch?lang=zh-CN"
    fields = "status,message,country,regionName,city,isp,query"
    for i in range(0, len(public_ips), BATCH_SIZE):
        chunk = public_ips[i: i + BATCH_SIZE]
        payload = [{"query": ip, "fields": fields} for ip in chunk]
        try:
            resp = requests.post(url, json=payload, timeout=30)
            resp.raise_for_status()
            for item in resp.json():
                ip_addr = item.get("query", "")
                results[ip_addr] = format_online_result(item)
        except Exception as e:
            _log(f"在线查询失败: {e}")
        if i + BATCH_SIZE < len(public_ips):
            time.sleep(BATCH_INTERVAL)
    return results

def query_all_ips(ips):
    results = {}
    for ip in ips:
        local = local_ip_label(ip)
        if local:
            results[ip] = (local, local, local)
    pending = [ip for ip in ips if ip not in results]
    if not pending:
        return results
    searcher = get_offline_searcher()
    if searcher:
        _log(f"使用本地 ip2region 离线库查询 {len(pending)} 个IP...")
        results.update(query_offline(searcher, pending))
        found = sum(1 for v in results.values() if v[0] not in ("未知", "查询失败"))
        _log(f"离线查询完成，成功 {found}/{len(pending)}")
        return results
    _log(f"使用在线 API 批量查询 {len(pending)} 个IP...")
    online = query_online_batch(pending)
    for ip in pending:
        location = online.get(ip, "未知")
        parts = location.split()
        stat_province = parts[1] if len(parts) >= 2 and parts[0] == "中国" else parts[0] if parts else "未知"
        stat_city = parts[2] if len(parts) >= 3 else stat_province
        results[ip] = (location, stat_province, stat_city)
    return results

def extract_source_ips(file_path):
    import pandas as pd
    df = pd.read_excel(file_path)
    df.columns = df.columns.str.strip()
    col_names = list(df.columns)
    src_ip_col = None
    dst_ip_col = None
    for col in col_names:
        col_lower = col.lower() if isinstance(col, str) else ''
        if '源' in col and 'ip' in col_lower:
            src_ip_col = col
        elif '目' in col and 'ip' in col_lower:
            dst_ip_col = col
    if not src_ip_col or not dst_ip_col:
        _log(f"[!] 无法找到必需的列，文件: {file_path}")
        return [], [], []
    df = df.rename(columns={src_ip_col: '源 IP', dst_ip_col: '目标 IP'})
    # 分离已排除IP（本地公网IP）和待分析IP
    excluded_mask = df['源 IP'].apply(is_excluded_ip)
    excluded_df = df[excluded_mask].copy()
    df = df[~excluded_mask]
    # 收集被排除的源IP（去重）
    excluded_ips = excluded_df['源 IP'].drop_duplicates().tolist() if len(excluded_df) > 0 else []
    def get_ip_type(ip_str):
        try:
            ip = ipaddress.ip_address(str(ip_str).strip())
            return 'internal' if ip.is_private else 'external'
        except ValueError:
            return 'invalid'
    df['源IP类型'] = df['源 IP'].apply(get_ip_type)
    df['目标IP类型'] = df['目标 IP'].apply(get_ip_type)
    external_to_internal = df[(df['源IP类型'] == 'external') & (df['目标IP类型'] == 'internal')].copy()
    internal_df = df[df['源IP类型'] == 'internal'].copy()
    external_ips = []
    if len(external_to_internal) > 0:
        external_ips = external_to_internal['源 IP'].drop_duplicates().tolist()
    internal_ips = []
    if len(internal_df) > 0:
        internal_ips = internal_df['源 IP'].drop_duplicates().tolist()
    return external_ips, internal_ips, excluded_ips

def generate_ip_report(files, date):
    all_external_ips = set()
    all_internal_ips = set()
    all_excluded_ips = set()
    for f in files:
        _log(f'[+] 处理文件: {f.name}')
        file_path = str(f)
        external_ips, internal_ips, excluded_ips = extract_source_ips(file_path)
        all_external_ips.update(external_ips)
        all_internal_ips.update(internal_ips)
        all_excluded_ips.update(excluded_ips)
    _log(f'[+] 外网攻击IP去重后共 {len(all_external_ips)} 个，开始查询归属地...')
    _log(f'[+] 内网IP去重后共 {len(all_internal_ips)} 个，开始查询归属地...')
    _log(f'[+] 本地公网IP(已排除)去重后共 {len(all_excluded_ips)} 个')
    external_ip_list = list(all_external_ips)
    internal_ip_list = list(all_internal_ips)
    excluded_ip_list = list(all_excluded_ips)
    location_map = query_all_ips(external_ip_list)
    # 查询排除IP的归属地
    excluded_location_map = query_all_ips(excluded_ip_list) if excluded_ip_list else {}
    load_terminal_ip_table()
    # 加载 local_geos 配置用于标记本地IP
    try:
        conf = load_config()
        local_geos = conf.get('geos', set())
    except Exception as e:
        _log(f'[!] 加载配置失败，local_geos 回退为空: {e}')
        local_geos = set()
    # 使用 runtime_dir（脚本/exe所在目录）作为输出基础目录
    out_dir = Path(runtime_dir) / 'output'
    out_dir.mkdir(exist_ok=True)
    output_file = out_dir / f'IP归属分析结果-{date}.xlsx'
    
    # 懒加载openpyxl模块
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
    
    wb = Workbook()
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                         top=Side(style='thin'), bottom=Side(style='thin'))
    # 本地IP高亮填充
    local_fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
    
    ws1 = wb.active
    ws1.title = '外网攻击IP归属'
    headers1 = ['序号', 'IP地址', '归属地', '备注']
    ws1.append(headers1)
    for cell in ws1[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border
    for idx, ip in enumerate(external_ip_list, start=1):
        location, _, _ = location_map.get(ip, ("未知", "未知", "未知"))
        # 检查是否为本地IP
        remark = ''
        is_local = False
        if local_geos and any(g in str(location) for g in local_geos):
            remark = '本地IP'
            is_local = True
        ws1.append([idx, ip, location, remark])
        for col in range(1, 5):
            ws1.cell(row=idx + 1, column=col).border = thin_border
            if is_local:
                ws1.cell(row=idx + 1, column=col).fill = local_fill
        ws1.cell(row=idx + 1, column=1).alignment = Alignment(horizontal='center', vertical='center')
        ws1.cell(row=idx + 1, column=2).alignment = Alignment(horizontal='left', vertical='center')
        ws1.cell(row=idx + 1, column=3).alignment = Alignment(wrap_text=True, vertical='center')
        ws1.cell(row=idx + 1, column=4).alignment = Alignment(horizontal='center', vertical='center')
    ws1.column_dimensions["A"].width = 8
    ws1.column_dimensions["B"].width = 18
    ws1.column_dimensions["C"].width = 40
    ws1.column_dimensions["D"].width = 12
    ws2 = wb.create_sheet('内网IP归属')
    headers2 = ['序号', 'IP地址', '归属地']
    ws2.append(headers2)
    for cell in ws2[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border
    for idx, ip in enumerate(internal_ip_list, start=1):
        location = get_terminal_location(ip)
        ws2.append([idx, ip, location])
        for col in range(1, 4):
            ws2.cell(row=idx + 1, column=col).border = thin_border
        ws2.cell(row=idx + 1, column=1).alignment = Alignment(horizontal='center', vertical='center')
        ws2.cell(row=idx + 1, column=2).alignment = Alignment(horizontal='left', vertical='center')
        ws2.cell(row=idx + 1, column=3).alignment = Alignment(wrap_text=True, vertical='center')
    ws2.column_dimensions["A"].width = 8
    ws2.column_dimensions["B"].width = 18
    ws2.column_dimensions["C"].width = 40
    # 第三 sheet：本地公网IP（已排除）
    if excluded_ip_list:
        ws3 = wb.create_sheet('本地公网IP(已排除)')
        headers3 = ['序号', 'IP地址', '归属地', '类型']
        ws3.append(headers3)
        for cell in ws3[1]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = thin_border
        for idx, ip in enumerate(excluded_ip_list, start=1):
            location, _, _ = excluded_location_map.get(ip, ("未知", "未知", "未知"))
            label = EXCLUDED_IP_LABELS.get(ip, '业务IP')
            ws3.append([idx, ip, location, label])
            for col in range(1, 5):
                ws3.cell(row=idx + 1, column=col).border = thin_border
                ws3.cell(row=idx + 1, column=col).fill = local_fill
            ws3.cell(row=idx + 1, column=1).alignment = Alignment(horizontal='center', vertical='center')
            ws3.cell(row=idx + 1, column=2).alignment = Alignment(horizontal='left', vertical='center')
            ws3.cell(row=idx + 1, column=3).alignment = Alignment(wrap_text=True, vertical='center')
            ws3.cell(row=idx + 1, column=4).alignment = Alignment(horizontal='center', vertical='center')
        ws3.column_dimensions["A"].width = 8
        ws3.column_dimensions["B"].width = 18
        ws3.column_dimensions["C"].width = 40
        ws3.column_dimensions["D"].width = 22
    try:
        wb.save(output_file)
    except PermissionError:
        # 文件被占用（可能在Excel中打开），使用带时间戳的备用文件名
        alt_file = out_dir / f'IP归属分析结果-{date}-{datetime.now().strftime("%H%M%S")}.xlsx'
        wb.save(alt_file)
        _log(f'[!] 原文件被占用，已保存为: {alt_file}')
        output_file = alt_file
    _log(f'[✓] IP归属分析完成，结果已保存至: {output_file}')
    return output_file

def _find_file(filename):
    """查找文件，优先从exe目录，其次从临时解压目录"""
    # 先从exe目录查找
    exe_path = os.path.join(runtime_dir, filename)
    if os.path.exists(exe_path):
        return exe_path
    # 如果有临时解压目录，从那里查找
    if meipass_dir:
        meipass_path = os.path.join(meipass_dir, filename)
        if os.path.exists(meipass_path):
            return meipass_path
    return exe_path  # 返回exe目录路径（可能不存在）

def load_config():
    cfg = configparser.ConfigParser()
    # 优先从exe目录读取，其次从临时解压目录
    config_path = _find_file('config.ini')
    cfg.read(config_path, encoding='utf-8')
    conf = {}
    conf['title'] = cfg['base'].get('report_title', fallback='网络安全值守保障日报')
    conf['pattern'] = cfg['base'].get('input_pattern', fallback='*.xlsx')
    conf['out_dir'] = cfg['base'].get('output_dir', fallback='output')
    conf['intel_file'] = cfg['base'].get('intel_file', fallback='intel.csv')
    ranges_raw = cfg['network'].get('ranges', fallback='')
    nets = [ipaddress.ip_network(x.strip()) for x in ranges_raw.splitlines()
            if x.strip() and not x.strip().startswith('#')]
    terminal_ips = set()
    # 复用 load_terminal_ip_table()（已缓存，避免重复读取同一文件）
    terminal_table = load_terminal_ip_table()
    if terminal_table:
        terminal_ips = set(terminal_table.keys())
        terminal_nets = list(ipaddress.collapse_addresses([ipaddress.ip_address(ip) for ip in terminal_ips]))
        nets.extend(terminal_nets)
        _log(f'[+] 加载终端IP地址表: {len(terminal_ips)}个IP, 合并为{len(terminal_nets)}个网段')
    conf['nets'] = nets
    conf['zones'] = {x.strip() for x in cfg['network'].get('internal_zones', '').split(',') if x.strip()}
    conf['geos'] = {x.strip() for x in cfg['network'].get('local_geos', '').split(',') if x.strip()}
    probes_raw = cfg['health'].get('probes', fallback='')
    conf['probes'] = [tuple(x.strip().split('|')) for x in probes_raw.splitlines()
                      if x.strip() and not x.strip().startswith('#') and '|' in x]
    conf['retention'] = cfg['health'].getint('log_retention_days', fallback=180)
    conf['top'] = cfg['report'].getint('top_events', fallback=5)
    conf['crit_levels'] = {x.strip() for x in cfg['report'].get('critical_levels', '严重,高危').split(',') if x.strip()}
    conf['ban_levels'] = {x.strip() for x in cfg['report'].get('ban_levels', '高危,严重').split(',') if x.strip()}
    return conf

def classify(ip, zone, geo, conf):
    try:
        addr = ipaddress.ip_address(str(ip).strip())
    except ValueError:
        return '未知'
    for net in conf['nets']:
        if addr in net:
            return '内网'
    # 使用None检查替代pd.isna
    zone = '' if zone is None or (hasattr(zone, '__class__') and zone.__class__.__name__ == 'NaNType') else str(zone)
    geo = '' if geo is None or (hasattr(geo, '__class__') and geo.__class__.__name__ == 'NaNType') else str(geo)
    if zone in conf['zones'] and any(g in geo for g in conf['geos']):
        return '内网'
    return '外网' if addr.is_global else '待确认'

def load_single_file(path, conf):
    import pandas as pd
    df = pd.read_excel(path, sheet_name=0, dtype={'源 IP': str, '目的 IP': str})
    df.columns = df.columns.str.strip()
    cols = []
    seen = set()
    for c in df.columns:
        if c in seen:
            i = 1
            while f'{c}_{i}' in seen:
                i += 1
            cols.append(f'{c}_{i}')
            seen.add(f'{c}_{i}')
        else:
            cols.append(c)
            seen.add(c)
    df.columns = cols
    col_map = {}
    target_counts = {}
    for c in df.columns:
        base = c.split('_')[0] if '_' in c else c
        target = None
        if '源' in base and 'IP' in base and '目的' not in base:
            target = '源IP'
        elif '目的' in base and 'IP' in base:
            target = '目的IP'
        elif '攻击' in base and ('名称' in base or '类型' in base):
            target = '攻击名称'
        elif '威胁' in base and '等级' in base:
            target = '威胁等级'
        elif '源' in base and '区域' in base:
            target = '源区域'
        elif '源' in base and '地理' in base:
            target = '源地理信息'
        elif '情报' in base or 'IOC' in base.upper():
            target = '情报IOC'
        elif '攻击' in base and '阶段' in base:
            target = '攻击阶段'
        elif '攻击' in base and '状态' in base:
            target = '攻击状态'
        if target:
            if target in target_counts:
                target_counts[target] += 1
                col_map[c] = f'{target}_{target_counts[target]}'
            else:
                target_counts[target] = 1
                col_map[c] = target
    df = df.rename(columns=col_map)
    for must in ['源IP', '目的IP', '攻击名称', '威胁等级']:
        if must not in df.columns:
            df[must] = ''
    df['网络类型'] = df.apply(
        lambda r: classify(r['源IP'], r.get('源区域', ''), r.get('源地理信息', ''), conf), axis=1)
    return df

def load_and_classify(paths, conf):
    dfs = []
    import pandas as pd
    for path in paths:
        _log(f'[+] 读取文件: {path.name}')
        df = load_single_file(path, conf)
        dfs.append(df)
    if not dfs:
        raise ValueError('没有读取到任何数据')
    df = pd.concat(dfs, ignore_index=True)
    _log(f'[+] 合并后共 {len(df)} 条记录')
    return df

LEVELS = ['严重', '高危', '中危', '低危']

def analyze(df):
    total = len(df)
    internal = df[df['网络类型'] == '内网']
    external = df[df['网络类型'] == '外网']
    def by_level(sub):
        return {lv: int((sub['威胁等级'] == lv).sum()) for lv in LEVELS}
    external_to_internal = external.copy()
    if '目的IP' in external.columns:
        external_to_internal = external[external['目的IP'].apply(is_private_ip)]
    external_to_internal = external_to_internal[~external_to_internal['源IP'].apply(is_excluded_ip)]
    ban_count = int(external_to_internal['源IP'].nunique()) if len(external_to_internal) > 0 else 0
    return {
        'total': total,
        'internal': internal,
        'external': external,
        'int_count': len(internal),
        'ext_count': len(external),
        'int_level': by_level(internal),
        'ext_level': by_level(external),
        'ban_count': ban_count,
    }

def _set_run_font(run, size=None, bold=None):
    """统一设置 run 字体：英文 Times New Roman，中文宋体"""
    from docx.oxml.ns import qn
    from docx.shared import Pt
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold

def _parse_lines(text, skip_example=True):
    """解析多行文本：按换行分割，去除空行和序号前缀"""
    if not text or not text.strip():
        return []
    items = []
    for line in text.split('\n'):
        line = line.strip()
        if line and (not skip_example or not line.startswith('示例：')):
            line = re.sub(r'^\d+\.\s*', '', line)
            items.append(line)
    return items

def _add_para(doc, text, bold=False, size=None):
    """添加段落并统一字体"""
    p = doc.add_paragraph()
    r = p.add_run(str(text))
    _set_run_font(r, size=size, bold=bold)
    return p

def _add_heading(doc, text, level=1):
    """添加标题并统一字体"""
    h = doc.add_heading(text, level)
    for run in h.runs:
        _set_run_font(run)
    return h

def _add_numbered_list(doc, items, start=1):
    """添加编号列表段落"""
    for i, item in enumerate(items, start):
        _add_para(doc, f'{i}. {item}')

def _add_table(doc, headers, widths, rows):
    """创建表格并填充数据，统一设置样式和列宽"""
    from docx.shared import Pt
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    _hdr(tbl, headers)
    for row_data in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            _set_cell(cells[i], val)
    _fit_table(tbl, [Pt(w) for w in widths])
    return tbl

def _hdr(table, headers):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(p.add_run(h), size=9, bold=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def _set_cell(cell, text, bold=False, size=9, padding=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    _set_run_font(p.add_run(str(text)), size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 仅当指定 padding 时才设置，否则使用 Word 默认值
    if padding is not None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = tcPr.find(qn('w:tcMar'))
        if tcMar is None:
            tcMar = OxmlElement('w:tcMar')
            tcPr.append(tcMar)
        for child in list(tcMar):
            tcMar.remove(child)
        for side, val in padding.items():
            elem = OxmlElement(f'w:{side}')
            elem.set(qn('w:w'), val)
            elem.set(qn('w:type'), 'dxa')
            tcMar.append(elem)

def _fit_table(table, widths=None):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    if widths is None:
        widths = []
    # 在所有行上设置单元格宽度
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]
    # 设置表格固定布局（与用户调整的文档一致）
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')
    # 允许跨页断行
    for row in table.rows:
        trPr = row._tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            row._tr.insert(0, trPr)
        cant_split = trPr.find(qn('w:cantSplit'))
        if cant_split is None:
            cant_split = OxmlElement('w:cantSplit')
            trPr.append(cant_split)

def render(conf, df, stats, health_rows, intel_list, date, out_path, work_summary=None, follow_items=None, intel_items=None):
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = docx.Document()
    # 统一设置样式字体
    def _apply_style_font(style):
        f = style.font
        f.name = 'Times New Roman'
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rFonts.set(qn(attr), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), '宋体')
    
    _apply_style_font(doc.styles['Normal'])
    doc.styles['Normal'].font.size = Pt(10.5)
    for name in ('Heading 1', 'Heading 2', 'Title'):
        try:
            _apply_style_font(doc.styles[name])
        except KeyError:
            pass
    h = _add_heading(doc, conf['title'], 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    r = sub.add_run(f'日期：{date[:4]}-{date[4:6]}-{date[6:]}    编制：网络安全值守组    密级：内部')
    _set_run_font(r, bold=True)
    sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 一、当日态势概览与重点工作总结
    _add_heading(doc, '一、当日态势概览与重点工作总结', 1)
    total = max(stats['total'], 1)
    auto_summary = (
        f"今日共捕获告警 {stats['total']} 起，其中内网 {stats['int_count']} 起"
        f"（{stats['int_count']/total*100:.1f}%），外网 {stats['ext_count']} 起"
        f"（{stats['ext_count']/total*100:.1f}%），累计封禁 IP {stats['ban_count']} 个。"
    )
    _add_para(doc, f'1. {auto_summary}')
    work_items = _parse_lines(work_summary)
    _add_numbered_list(doc, work_items, start=2)
    
    # 二、关键指标
    _add_heading(doc, '二、关键指标', 1)
    _add_table(doc, ['内网告警', '次数', '外网告警', '次数'], [110, 111, 110, 111],
        [(lv, stats['int_level'].get(lv, 0), lv, stats['ext_level'].get(lv, 0)) for lv in LEVELS],
    )
    for idx, (title, body) in enumerate([
        ('可疑域名解析请求分析', '通过定期对应用进行安全事件监测，暂未发现可疑域名解析请求事件。'),
        ('敏感文件分析', '通过定期对应用进行安全事件监测，暂未发现敏感文件。'),
        ('黑链挂马分析', '通过定期对应用进行安全事件监测，暂未发现黑链/挂马事件。'),
        ('可用性分析', '通过定期对应用进行安全事件监测，暂未发现存在可用性事件。'),
        ('篡改分析', '通过定期对应用进行安全事件监测，暂未发现网页篡改事件。'),
        ('域名劫持分析', '通过定期对应用进行安全事件监测,暂未发现域名劫持事件。'),
    ], start=1):
        _add_heading(doc, f'{idx}. {title}', 2)
        _add_para(doc, body)
    
    # 三、研判与处置流程
    _add_heading(doc, '三、研判与处置流程', 1)
    for s in [
        '流量采集：安全设备日志导出 + NDR 探针 + 防火墙 IPS',
        '情报比对：威胁情报平台 IOC 命中 / ASN 与归属地归类',
        '内外网判定：源 IP 静态段 + 资产归属 + 流量方向三重校验',
        '分级处置：P0 即时处置 / P1 当日处置 / P2 观察 / P3 沉淀',
    ]:
        p = doc.add_paragraph(style='List Number')
        _set_run_font(p.add_run(s))
    
    # 四、资产健康检查
    _add_heading(doc, '四、资产健康检查', 1)
    _add_para(doc, '值守期间持续监控安全设备运行状态，确保监测能力不降级。')
    _add_table(doc, ['探针名称', 'IP', '状态', '备注'], [120, 100, 111, 111],
        [(name, ip, '正常', '') for name, ip, *_ in health_rows])
    _add_para(doc, f'日志留存：{conf["retention"]} 天。')
    
    # 五、外网攻击研判与处置
    _add_heading(doc, '五、外网攻击研判与处置', 1)
    ext = stats['external']
    if len(ext) > 0:
        grp = ext.groupby('攻击名称', sort=False).agg({'威胁等级': 'first', '源IP': 'count'}).reset_index()
        ext_rows = [(idx, row['攻击名称'], f"{int(row['源IP'])} 起", '流量特征+情报比对', '已封禁')
                    for idx, (_, row) in enumerate(grp.iterrows(), start=1)]
    else:
        ext_rows = [('（无外网告警）', '', '', '', '')]
    _add_table(doc, ['序号', '威胁类型', '命中次数', '研判依据', '状态'], [41, 203, 51, 106, 41], ext_rows)
    
    # 六、内网异常研判与处置
    _add_heading(doc, '六、内网异常研判与处置', 1)
    intdf = stats['internal']
    if len(intdf) > 0:
        agg_map = {'源IP': 'count'}
        if '威胁等级' in intdf.columns:
            agg_map['威胁等级'] = 'first'
        grp = intdf.groupby('攻击名称', sort=False).agg(agg_map).reset_index()
        int_rows = []
        for idx, (_, row) in enumerate(grp.iterrows(), start=1):
            lvl = row.get('威胁等级', '') if '威胁等级' in row else ''
            status = '取证/核查' if lvl in conf['crit_levels'] else '观察'
            int_rows.append((idx, row['攻击名称'], f"{int(row['源IP'])} 起", '流量特征+情报比对', status))
    else:
        int_rows = [('（无内网告警）', '', '', '', '')]
    _add_table(doc, ['序号', '威胁类型', '命中次数', '研判依据', '状态'], [41, 203, 51, 106, 41], int_rows)
    
    # 七、重点事件研判
    _add_heading(doc, '七、重点事件研判', 1)
    key = df[df['威胁等级'].isin(conf['crit_levels'])].copy()
    if len(key) > 0:
        key['_p'] = key['威胁等级'].map({lv: i for i, lv in enumerate(LEVELS)})
        key = key.sort_values('_p')
        grp = key.groupby(['攻击名称', '源IP'], sort=False).size().reset_index(name='count')
        grp = grp.sort_values('count', ascending=False).head(conf['top'])
        evt_name_map = {}
        for idx, row in enumerate(grp[['攻击名称']].drop_duplicates().values, 1):
            evt_name_map[row[0]] = idx
        key_rows = [(evt_name_map[row['攻击名称']], row['攻击名称'], row['源IP'], row['count'])
                    for _, row in grp.iterrows()]
        _add_table(doc, ['序号', '事件名称', '源地址', '攻击次数'], [41, 145, 135, 121], key_rows)
    else:
        _add_para(doc, '今日无严重/高危级事件。')
    
    # 八、情报动态
    _add_heading(doc, '八、情报动态', 1)
    _add_para(doc, '当日需关注的新增 CVE / 行业预警：')
    intel_parsed = _parse_lines(intel_items, skip_example=False)
    if intel_parsed:
        _add_numbered_list(doc, intel_parsed)
    elif intel_list:
        _add_table(doc, ['类型', '编号', '风险', '关联资产', '应对/时限'], [60, 85, 65, 115, 117],
            [(item.get('类型', ''), item.get('编号', ''), item.get('风险', ''),
              item.get('关联资产', ''), f"{item.get('应对','')} / {item.get('时限','')}")
             for item in intel_list])
    else:
        _add_para(doc, '（暂无新增情报，详见威胁情报平台）')
    
    # 九、待跟进事项
    _add_heading(doc, '九、待跟进事项', 1)
    default_items = [
        '内网高危告警溯源与处置闭环',
        '外网封禁 IP 清单同步至边界防火墙',
        '失陷终端取证与隔离',
        '资产健康检查异常处理',
    ]
    follow_parsed = _parse_lines(follow_items, skip_example=False)
    _add_numbered_list(doc, follow_parsed if follow_parsed else default_items)
    
    doc.save(out_path)

def load_intel(conf):
    p = Path(os.path.join(runtime_dir, conf['intel_file']))
    if not p.exists():
        return []
    try:
        import pandas as pd
        return pd.read_csv(p).fillna('').to_dict('records')
    except Exception as e:
        _log(f'[!] 加载情报文件失败: {e}')
        return []

def pick_input_and_date(pattern):
    # 文件名排除关键词（不区分大小写）：非安全告警类的辅助/输出/配置文件
    _EXCLUDE_KEYS = (
        '终端ip地址表', '终端ip', '业务ip', '业务IP', 'biz_ip',
        'ip归属分析', 'IP归属分析', 'ip分析', 'IP分析',
        '值守保障日报', '值守日报', '日报', '日报汇总',
        'config', '配置', 'config.ini',
        'version', '情报', 'intel',
    )

    def _is_excluded(fpath):
        s = fpath.stem
        for k in _EXCLUDE_KEYS:
            if k.lower() in s.lower():
                return k
        return None

    files = sorted(Path(runtime_dir).glob(pattern), key=os.path.getmtime, reverse=True)
    if not files:
        files = sorted(Path(runtime_dir).glob('*.xls*'), key=os.path.getmtime, reverse=True)

    # 先输出候选与排除情况，便于定位
    excluded_info = []
    filtered = []
    for f in files:
        reason = _is_excluded(f)
        if reason:
            excluded_info.append(f'  ! {f.name}  已排除(命中关键词: {reason})')
        else:
            filtered.append(f)
    files = filtered
    if excluded_info:
        _log('[*] 自动识别排除文件:')
        _log('\n'.join(excluded_info))

    if not files:
        raise FileNotFoundError('未找到安全告警类 Excel，请把安全设备导出的日志放到脚本目录，'
                                '并确保文件名不含「业务IP/终端IP/值守日报/IP归属分析」等排除关键词')

    # 锚点优先选择: stem 中包含 8 位日期 的最新 mtime 文件（真实告警导出一般带时间戳）
    dated_files = [f for f in files if re.search(r'(\d{8})', f.stem)]
    if dated_files:
        anchor = sorted(dated_files, key=os.path.getmtime, reverse=True)[0]
    else:
        anchor = files[0]  # 退化: 所有文件都没带日期，取 mtime 最新

    stem = anchor.stem
    m = re.search(r'(\d{8})', stem)
    if m:
        date = m.group(1)
        _log(f'[*] 自动识别锚点文件: {anchor.name} (日期: {date})')
    else:
        date = datetime.now().strftime('%Y%m%d')
        _log(f'[*] 自动识别锚点文件: {anchor.name} (文件名无日期，使用今日: {date})')

    target_files = []
    for f in files:
        fm = re.search(r'(\d{8})', f.stem)
        if fm and fm.group(1) == date:
            target_files.append(f)
    if not target_files:
        target_files = [anchor]
    _log(f'[*] 自动识别结果: 共 {len(target_files)} 个文件 (日期: {date})')
    return target_files, date

def generate_daily_report(files, date, work_summary=None, follow_items=None, intel_items=None):
    conf = load_config()
    df = load_and_classify(files, conf)
    stats = analyze(df)
    _log(f'[+] 总告警 {stats["total"]} | 内网 {stats["int_count"]} | 外网 {stats["ext_count"]} | 封禁 {stats["ban_count"]}')
    health_rows = conf['probes']
    intel_list = load_intel(conf)
    # 使用 runtime_dir（脚本/exe所在目录）作为输出基础目录
    out_dir = Path(runtime_dir) / conf['out_dir']
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / f'值守保障日报{date}.docx'
    render(conf, df, stats, health_rows, intel_list, date, str(out_path), work_summary, follow_items, intel_items)
    _log(f'[✓] 值守日报已生成: {out_path}')
    return out_path

def cli_main():
    """命令行模式入口"""
    print('=' * 60)
    print('网络安全值守保障日报整合脚本')
    print('功能：1. IP归属分析  2. 值守日报生成')
    print('=' * 60)
    input_files, date = pick_input_and_date('*.xlsx')
    print(f'\n[+] 日期: {date}')
    print(f'[+] 发现 {len(input_files)} 个待处理文件:')
    for f in input_files:
        print(f'    - {f.name}')
    print('\n' + '=' * 60)
    print('步骤1: IP归属分析')
    print('=' * 60)
    ip_report_path = generate_ip_report(input_files, date)
    print('\n' + '=' * 60)
    print('步骤2: 生成值守日报')
    print('=' * 60)
    daily_report_path = generate_daily_report(input_files, date)
    print('\n' + '=' * 60)
    print('全部完成！')
    print(f'IP归属分析结果: {ip_report_path}')
    print(f'值守日报: {daily_report_path}')
    print('=' * 60)

# ==================== GUI 界面部分 ====================

# GUI 文本框示例占位文本
EXAMPLE_WORK_SUMMARY = '示例：\n1. 完成防火墙规则优化\n2. 处置高危漏洞告警'
EXAMPLE_INTEL_ITEMS = '示例：\n1. CVE-2024-XXXX 高危漏洞，需尽快修复\n2. 新发现XX行业专项攻击预警'

class DailyReportGUI:
    def __init__(self, master):
        self.master = master
        master.title('网络安全值守保障日报')
        master.geometry('850x750')
        master.resizable(False, False)
        
        self.input_files = []
        self.date_var = StringVar(value=datetime.now().strftime('%Y%m%d'))
        
        # 注册全局日志回调
        global _gui_log_callback, _gui_progress_callback
        _gui_log_callback = self._log
        _gui_progress_callback = self._set_progress
        
        self._build_ui()

        # 启动时后台自动检查更新（静默模式，不打扰用户）
        threading.Thread(target=self._check_update_startup, daemon=True).start()

    def _build_ui(self):
        # 顶部标题
        title_frame = Frame(self.master, padx=10, pady=5)
        title_frame.pack(fill='x')
        Label(title_frame, text='网络安全值守保障日报', font=('宋体', 16, 'bold')).pack(side='left')
        Label(title_frame, text=f'v{APP_VERSION}', font=('宋体', 9), fg='#666666').pack(side='left', padx=(10, 0))
        Button(title_frame, text='检查更新', command=self._check_update_manual, width=10,
               font=('宋体', 9)).pack(side='right')
        
        # 文件选择区域
        file_frame = Frame(self.master, padx=10, pady=5)
        file_frame.pack(fill='x')
        
        Label(file_frame, text='输入文件:', font=('宋体', 10)).pack(anchor='w')
        
        btn_frame = Frame(file_frame)
        btn_frame.pack(fill='x', pady=2)
        
        Button(btn_frame, text='选择文件', command=self._select_files, width=15,
               font=('宋体', 10)).pack(side='left')
        Button(btn_frame, text='自动识别', command=self._auto_detect, width=15,
               font=('宋体', 10)).pack(side='left', padx=5)
        Button(btn_frame, text='清空列表', command=self._clear_files, width=15,
               font=('宋体', 10)).pack(side='right')
        
        # 文件列表
        list_frame = Frame(file_frame)
        list_frame.pack(fill='x', pady=2)
        
        scrollbar = Scrollbar(list_frame, orient='vertical')
        self.file_listbox = Listbox(list_frame, yscrollcommand=scrollbar.set,
                                    font=('宋体', 9), selectmode='extended', height=5)
        scrollbar.config(command=self.file_listbox.yview)
        scrollbar.pack(side='right', fill='y')
        self.file_listbox.pack(side='left', fill='x', expand=True)
        
        # IP导入区域（业务IP + 终端IP表 并列一行）
        ip_import_frame = Frame(self.master, padx=10, pady=3)
        ip_import_frame.pack(fill='x')

        self.excluded_ip_label_var = StringVar(value='未导入')
        Button(ip_import_frame, text='导入业务IP', command=self._load_biz_ips, width=12,
               font=('宋体', 10)).pack(side='left')
        Label(ip_import_frame, textvariable=self.excluded_ip_label_var,
              font=('宋体', 9), fg='gray').pack(side='left', padx=(3, 15))

        self.terminal_ip_label_var = StringVar(value='未导入')
        Button(ip_import_frame, text='导入终端IP表', command=self._load_terminal_ips, width=14,
               font=('宋体', 10)).pack(side='left')
        Label(ip_import_frame, textvariable=self.terminal_ip_label_var,
              font=('宋体', 9), fg='gray').pack(side='left', padx=5)
        
        # 日期输入
        date_frame = Frame(self.master, padx=10, pady=3)
        date_frame.pack(fill='x')
        
        Label(date_frame, text='日期:', font=('宋体', 10)).pack(side='left')
        Entry(date_frame, textvariable=self.date_var, width=12,
              font=('宋体', 10)).pack(side='left', padx=5)
        
        # 重点工作总结输入区域
        work_frame = Frame(self.master, padx=10, pady=3)
        work_frame.pack(fill='x')
        
        Label(work_frame, text='重点工作总结（每行一项）:', font=('宋体', 10)).pack(anchor='w')
        work_inner = Frame(work_frame)
        work_inner.pack(fill='x')
        work_scroll = Scrollbar(work_inner, orient='vertical')
        self.work_summary_text = Text(work_inner, font=('宋体', 10),
                                      yscrollcommand=work_scroll.set,
                                      height=3, wrap='word')
        work_scroll.config(command=self.work_summary_text.yview)
        work_scroll.pack(side='right', fill='y')
        self.work_summary_text.pack(side='left', fill='x', expand=True)
        self.work_summary_text.insert(END, '示例：\n1. 完成防火墙规则优化\n2. 处置高危漏洞告警')
        
        # 待跟进事项输入区域
        follow_frame = Frame(self.master, padx=10, pady=3)
        follow_frame.pack(fill='x')
        
        Label(follow_frame, text='待跟进事项（每行一项，不填则使用默认）:', font=('宋体', 10)).pack(anchor='w')
        follow_inner = Frame(follow_frame)
        follow_inner.pack(fill='x')
        follow_scroll = Scrollbar(follow_inner, orient='vertical')
        self.follow_items_text = Text(follow_inner, font=('宋体', 10),
                                       yscrollcommand=follow_scroll.set,
                                       height=3, wrap='word')
        follow_scroll.config(command=self.follow_items_text.yview)
        follow_scroll.pack(side='right', fill='y')
        self.follow_items_text.pack(side='left', fill='x', expand=True)
        
        # 情报动态输入区域
        intel_frame = Frame(self.master, padx=10, pady=3)
        intel_frame.pack(fill='x')
        
        Label(intel_frame, text='情报动态（每行一项，不填则使用默认表格）:', font=('宋体', 10)).pack(anchor='w')
        intel_inner = Frame(intel_frame)
        intel_inner.pack(fill='x')
        intel_scroll = Scrollbar(intel_inner, orient='vertical')
        self.intel_items_text = Text(intel_inner, font=('宋体', 10),
                                      yscrollcommand=intel_scroll.set,
                                      height=3, wrap='word')
        intel_scroll.config(command=self.intel_items_text.yview)
        intel_scroll.pack(side='right', fill='y')
        self.intel_items_text.pack(side='left', fill='x', expand=True)
        self.intel_items_text.insert(END, '示例：\n1. CVE-2024-XXXX 高危漏洞，需尽快修复\n2. 新发现XX行业专项攻击预警')
        
        # 日志输出区域
        log_frame = Frame(self.master, padx=10, pady=2)
        log_frame.pack(fill='x')
        
        Label(log_frame, text='执行日志:', font=('宋体', 10)).pack(anchor='w')
        
        log_inner = Frame(log_frame, height=100)
        log_inner.pack(fill='x', pady=2)
        log_inner.pack_propagate(False)
        
        self.log_text = Text(log_inner, font=('Consolas', 9), state=DISABLED,
                             wrap=NONE, bg='#f5f5f5')
        log_scroll_y = Scrollbar(log_inner, orient='vertical', command=self.log_text.yview)
        log_scroll_x = Scrollbar(log_inner, orient='horizontal', command=self.log_text.xview)
        self.log_text.configure(yscrollcommand=log_scroll_y.set, xscrollcommand=log_scroll_x.set)
        
        log_scroll_y.pack(side='right', fill='y')
        log_scroll_x.pack(side='bottom', fill='x')
        self.log_text.pack(side='left', fill='both', expand=True)
        
        # 下载进度条
        progress_frame = Frame(self.master, padx=10, pady=2)
        progress_frame.pack(fill='x')
        self.progress_label = Label(progress_frame, text='', font=('宋体', 9), width=20, anchor='w')
        self.progress_label.pack(side='left')
        from tkinter.ttk import Progressbar
        self.progress_bar = Progressbar(progress_frame, mode='determinate', length=600)
        self.progress_bar.pack(side='left', fill='x', expand=True, padx=5)
        
        # 底部按钮区域
        btn_frame = Frame(self.master, padx=10, pady=5)
        btn_frame.pack(fill='x')
        
        self.run_btn = Button(btn_frame, text='开始生成', command=self._run,
                              font=('宋体', 12, 'bold'), width=20, bg='#4CAF50', fg='white')
        self.run_btn.pack(side='left', padx=5)
        
        self.open_ip_btn = Button(btn_frame, text='打开IP分析结果', command=self._open_ip_report,
                                  font=('宋体', 10), width=20, state=DISABLED)
        self.open_ip_btn.pack(side='left', padx=5)
        
        self.open_report_btn = Button(btn_frame, text='打开值守日报', command=self._open_daily_report,
                                      font=('宋体', 10), width=20, state=DISABLED)
        self.open_report_btn.pack(side='right', padx=5)
        
        # 结果路径
        self.ip_report_path = None
        self.daily_report_path = None
    
    def _log(self, msg):
        def do_log():
            self.log_text.config(state=NORMAL)
            timestamp = datetime.now().strftime('%H:%M:%S')
            self.log_text.insert(END, f'[{timestamp}] {msg}\n')
            self.log_text.see(END)
            self.log_text.config(state=DISABLED)
        self.master.after(0, do_log)
    
    def _set_progress(self, value, maximum=None):
        def do_progress():
            if maximum is None:
                self.progress_bar.config(mode='indeterminate')
                if value > 0:
                    self.progress_bar.step(1)
                else:
                    self.progress_bar.stop()
                    self.progress_bar.config(mode='determinate')
            elif maximum == 0:
                self.progress_bar.config(mode='determinate', maximum=100, value=0)
                self.progress_label.config(text='')
            else:
                self.progress_bar.config(mode='determinate', maximum=maximum, value=value)
                pct = int(value / maximum * 100) if maximum > 0 else 0
                mb_done = value / (1024 * 1024)
                mb_total = maximum / (1024 * 1024)
                self.progress_label.config(text=f'{mb_done:.1f}/{mb_total:.1f}MB ({pct}%)')
        self.master.after(0, do_progress)
    
    def _load_biz_ips(self):
        f = filedialog.askopenfilename(
            title='选择业务IP Excel文件',
            filetypes=[('Excel文件', '*.xlsx *.xls'), ('所有文件', '*.*')],
            initialdir=runtime_dir
        )
        if not f:
            return
        count = load_external_excluded_ips(f)
        if count > 0:
            self.excluded_ip_label_var.set(f'已导入 {count} 条业务IP')
            self._log(f'已导入业务IP: {count} 条 (来源: {os.path.basename(f)})')
        else:
            self.excluded_ip_label_var.set('导入失败')
            messagebox.showwarning('提示', '未能从文件中加载到有效的IP，请检查文件格式（需包含IP列）')
            self._log(f'业务IP导入失败: {f}')

    def _load_terminal_ips(self):
        """导入终端IP地址表（外部Excel文件）"""
        f = filedialog.askopenfilename(
            title='选择终端IP地址表 Excel文件',
            filetypes=[('Excel文件', '*.xlsx *.xls'), ('所有文件', '*.*')],
            initialdir=runtime_dir
        )
        if not f:
            return
        try:
            set_terminal_ip_table_path(f)
            table = load_terminal_ip_table()
            count = len(table)
            if count > 0:
                self.terminal_ip_label_var.set(f'已导入 {count} 条终端IP')
                self._log(f'已导入终端IP表: {count} 条 (来源: {os.path.basename(f)})')
            else:
                self.terminal_ip_label_var.set('导入失败')
                messagebox.showwarning('提示', '未能从文件中加载到有效的IP，请检查文件格式')
                self._log(f'终端IP表导入失败: {f}')
        except Exception as e:
            self.terminal_ip_label_var.set('导入失败')
            self._log(f'终端IP表导入异常: {e}')

    def _select_files(self):
        files = filedialog.askopenfilenames(
            title='选择安全告警Excel文件',
            filetypes=[('Excel文件', '*.xlsx *.xls'), ('所有文件', '*.*')],
            initialdir=runtime_dir
        )
        if files:
            for f in files:
                if f not in self.input_files:
                    self.input_files.append(f)
                    self.file_listbox.insert(END, os.path.basename(f))
            # 从第一个文件名中提取日期
            first_file = files[0]
            stem = os.path.splitext(os.path.basename(first_file))[0]
            m = re.search(r'(\d{8})', stem)
            if m:
                self.date_var.set(m.group(1))
            self._log(f'已选择 {len(files)} 个文件')
    
    def _auto_detect(self):
        try:
            files, date = pick_input_and_date('*.xlsx')
            self.input_files = [str(f) for f in files]
            self.date_var.set(date)
            self.file_listbox.delete(0, END)
            for f in self.input_files:
                self.file_listbox.insert(END, os.path.basename(f))
            self._log(f'自动识别到 {len(files)} 个文件 (日期: {date})')
        except FileNotFoundError as e:
            messagebox.showwarning('提示', str(e))
            self._log(f'警告: {e}')
        except Exception as e:
            import traceback
            error_msg = f'_auto_detect回调错误: {str(e)}\n{traceback.format_exc()}'
            print(error_msg)
            messagebox.showerror('错误', error_msg)
    
    def _clear_files(self):
        # 获取选中的索引（从后往前删除，避免索引偏移）
        selected_indices = sorted(self.file_listbox.curselection(), reverse=True)
        if selected_indices:
            # 删除选中的文件
            for idx in selected_indices:
                del self.input_files[idx]
                self.file_listbox.delete(idx)
            self._log(f'已移除 {len(selected_indices)} 个选中文件')
        else:
            # 没有选中，清空全部
            self.input_files = []
            self.file_listbox.delete(0, END)
            self._log('文件列表已清空')
    
    def _run(self):
        try:
            if not self.input_files:
                messagebox.showwarning('提示', '请先选择或自动识别输入文件')
                return
            
            self.run_btn.config(state=DISABLED)
            self.open_ip_btn.config(state=DISABLED)
            self.open_report_btn.config(state=DISABLED)
            
            def worker():
                try:
                    date = self.date_var.get()
                    
                    # 获取用户输入的重点工作总结
                    work_summary = self.work_summary_text.get('1.0', END).strip()
                    
                    # 获取用户输入的待跟进事项
                    follow_items = self.follow_items_text.get('1.0', END).strip()
                    
                    # 获取用户输入的情报动态
                    intel_items = self.intel_items_text.get('1.0', END).strip()
                    # 如果输入内容等于示例文本，视为未填写
                    intel_default_example = EXAMPLE_INTEL_ITEMS
                    if intel_items == intel_default_example:
                        intel_items = None
                    
                    self._log('=' * 50)
                    self._log('步骤1: IP归属分析')
                    self._log('=' * 50)
                    
                    path_objects = [Path(f) for f in self.input_files]
                    self.ip_report_path = generate_ip_report(path_objects, date)
                    self._log(f'IP归属分析完成: {self.ip_report_path}')
                    
                    self._log('')
                    self._log('=' * 50)
                    self._log('步骤2: 生成值守日报')
                    self._log('=' * 50)
                    
                    self.daily_report_path = generate_daily_report(path_objects, date, work_summary, follow_items, intel_items)
                    self._log(f'值守日报生成完成: {self.daily_report_path}')
                    
                    self._log('')
                    self._log('=' * 50)
                    self._log('全部完成！')
                    self._log('=' * 50)
                    
                    self.master.after(0, self._on_complete)
                    
                except Exception as e:
                    self._log(f'错误: {str(e)}')
                    import traceback
                    self._log(traceback.format_exc())
                    self.master.after(0, lambda: self.run_btn.config(state=NORMAL))
            
            threading.Thread(target=worker, daemon=True).start()
        except Exception as e:
            import traceback
            error_msg = f'_run回调错误: {str(e)}\n{traceback.format_exc()}'
            print(error_msg)
            messagebox.showerror('错误', error_msg)
    
    def _on_complete(self):
        self.run_btn.config(state=NORMAL)
        self.open_ip_btn.config(state=NORMAL)
        self.open_report_btn.config(state=NORMAL)
        messagebox.showinfo('完成', 'IP归属分析和值守日报已生成')
    
    def _open_ip_report(self):
        if self.ip_report_path and os.path.exists(self.ip_report_path):
            os.startfile(str(self.ip_report_path))
        else:
            messagebox.showwarning('提示', 'IP分析结果文件不存在')
    
    def _open_daily_report(self):
        if self.daily_report_path and os.path.exists(self.daily_report_path):
            os.startfile(str(self.daily_report_path))
        else:
            messagebox.showwarning('提示', '值守日报文件不存在')

    # ---------------- 更新相关方法 ----------------
    def _check_update_manual(self):
        """手动检查更新按钮回调（强制弹窗）"""
        try:
            threading.Thread(target=self._run_update_with_gui, args=(True,), daemon=True).start()
        except Exception as e:
            import traceback
            self._log(f'检查更新启动失败: {e}\n{traceback.format_exc()}')

    def _check_update_startup(self):
        """启动时自动检查更新（静默模式：有更新才弹窗）"""
        try:
            self._run_update_with_gui(force_dialog=False)
        except Exception as e:
            # 启动自动检查失败不打扰用户，仅记录日志
            import traceback
            self._log(f'[更新] 自动检查失败: {e}\n{traceback.format_exc()}')

    def _run_update_with_gui(self, force_dialog):
        """带GUI回调的更新流程（线程内运行）"""
        def ask_confirm(msg):
            # 对话框必须在主线程弹出，通过after + 事件同步
            result = [False]
            done = threading.Event()
            def do_ask():
                try:
                    result[0] = messagebox.askyesno("发现更新", msg)
                finally:
                    done.set()
            self.master.after(0, do_ask)
            done.wait()
            return result[0]

        updater = AutoUpdater(
            progress_cb=self._set_progress,
            log_cb=self._log,
            ask_confirm_cb=ask_confirm,
        )
        updater.run_update_flow(force_dialog=force_dialog)

def update_worker_main(json_path):
    """
    更新 Worker 模式：不加载 GUI，纯文件操作。
    由 install_and_restart() 触发：主进程 os._exit 后，worker 副本负责
    备份旧 exe → 覆盖新 exe → 启动新程序。

    设计目的：
    - 避免 PowerShell 5.1 的安全校验（Security validation failure:
      parent process has different executable! / failed to obtain executable path...）
    - 避免 cmd/.bat 的中文路径 GBK/UTF-8 乱码
    - 100% 走 Python 文件 API（Copy-Item 级），中文路径零问题
    """
    import json, traceback, hashlib

    # ===== 解析参数 =====
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            params = json.load(f)
        parent_pid = int(params['parentPid'])
        old_exe    = str(params['oldExe'])
        new_exe    = str(params['newExe'])
        bak_exe    = str(params['bakExe'])
        work_dir   = str(params['workDir'])
        cleanups   = [str(params['jsonPath'])]
        if 'ps1Path' in params and params['ps1Path']: cleanups.append(str(params['ps1Path']))
    except Exception as e:
        # 尽量写日志（优先临时目录，避免 exe 目录无写权限）
        try:
            import tempfile as _tf
            log = os.path.join(_tf.gettempdir(), 'update_last.log')
            with open(log, 'a', encoding='utf-8') as f:
                f.write(f"[{datetime.now():%H:%M:%S}] 参数解析失败: {e}\n")
        except Exception:
            pass
        return 2

    # 日志双写：优先临时目录（保证有写权限），同时尝试 exe 同目录（方便用户找）
    import tempfile as _tf
    _tmp_log = os.path.join(_tf.gettempdir(), 'update_last.log')
    _exe_log = os.path.join(work_dir, 'update_last.log')
    log_file = _tmp_log

    def wlog(msg):
        ts = datetime.now().strftime('%H:%M:%S')
        # 双写：优先临时目录（一定有权限），同时尝试 exe 同目录（方便用户找）
        for _p in (_tmp_log, _exe_log):
            try:
                with open(_p, 'a', encoding='utf-8') as f:
                    f.write(f'[{ts}] {msg}\n')
            except Exception:
                pass

    wlog("=" * 40)
    wlog(f"开始更新 (pid={os.getpid()}) parentPid={parent_pid}")
    wlog(f"oldExe={old_exe}")
    wlog(f"newExe={new_exe}")
    wlog(f"exe大小={os.path.getsize(new_exe)}" if os.path.exists(new_exe) else "新exe不存在！")

    # ===== 检测父进程是否仍在（用 OpenProcess Win32，不触发任何安全校验） =====
    try:
        import ctypes
        from ctypes import wintypes
        k32 = ctypes.WinDLL('kernel32', use_last_error=True)
        PROCESS_QUERY_LIMITED_INFORMATION = 0x1000
        SYNCHRONIZE = 0x00100000
        def _is_alive(pid):
            h = k32.OpenProcess(SYNCHRONIZE | PROCESS_QUERY_LIMITED_INFORMATION, False, wintypes.DWORD(pid))
            if not h:
                return False
            # WaitForSingleObject 0 毫秒：若已退出立刻返回 WAIT_OBJECT_0
            WAIT_OBJECT_0 = 0
            WAIT_TIMEOUT = 0x00000102
            r = k32.WaitForSingleObject(h, 0)
            k32.CloseHandle(h)
            return r == WAIT_TIMEOUT
    except Exception as e:
        wlog(f"OpenProcess 不可用 ({e})，降级用 psutil/tasklist")
        def _is_alive(pid):
            # 兜底：tasklist /FO CSV 数字匹配
            try:
                import subprocess as _sp
                out = _sp.check_output(
                    ['tasklist', '/FI', f'PID eq {pid}', '/NH', '/FO', 'CSV'],
                    stderr=_sp.DEVNULL
                ).decode('gbk', errors='replace')
                return f'"{pid}"' in out
            except Exception:
                return False

    waited = 0
    while waited < 10:
        if not _is_alive(parent_pid):
            wlog(f"父进程已退出 (等待 {waited}s)")
            break
        time.sleep(1)
        waited += 1
    if waited >= 10:
        wlog(f"等待父进程超时 ({waited}s)，继续覆盖（有重试兜底）")
    time.sleep(0.3)

    # ===== 备份 =====
    if os.path.exists(old_exe):
        try:
            import shutil as _su
            _su.copy2(old_exe, bak_exe)
            wlog(f"已备份 -> {os.path.basename(bak_exe)}")
        except Exception as e:
            wlog(f"备份失败: {e}")

    # ===== 覆盖（重试 5 次，间隔 2s） =====
    ok = False
    last_err = None
    import shutil as _su
    for i in range(1, 6):
        try:
            _su.copy2(new_exe, old_exe)
            # 验证 MD5 一致
            def md5(p):
                h = hashlib.md5()
                with open(p, 'rb') as f:
                    for c in iter(lambda: f.read(1024*1024), b''):
                        h.update(c)
                return h.hexdigest().lower()
            if md5(old_exe) != md5(new_exe):
                raise RuntimeError("覆盖后 MD5 不一致")
            wlog(f"覆盖成功 ({i}/5)，MD5 校验通过")
            ok = True
            break
        except Exception as e:
            last_err = e
            wlog(f"覆盖失败 {i}/5: {e}")
            time.sleep(2)

    if not ok:
        wlog(f"覆盖失败，尝试回滚：{last_err}")
        if os.path.exists(bak_exe):
            try:
                _su.copy2(bak_exe, old_exe)
                wlog("回滚成功")
            except Exception as e:
                wlog(f"回滚失败: {e}")
        # 回滚后依旧失败，留日志返回
        return 1

    # ===== 启动新程序 =====
    started_ok = False
    try:
        DETACHED_PROCESS = 0x00000008
        CREATE_NEW_PROCESS_GROUP = 0x00000200
        CREATE_NO_WINDOW = 0x08000000
        flags = CREATE_NO_WINDOW | DETACHED_PROCESS | CREATE_NEW_PROCESS_GROUP
        # 注意：不使用 close_fds=True 与 creationflags 同时传时在老 Python 有兼容
        # 直接用 subprocess.Popen(executable, cwd=work_dir)，且不指定 stdio
        import subprocess as _sp
        _sp.Popen([old_exe], cwd=work_dir, creationflags=flags)
        wlog("已启动新程序 (Popen+DETACHED)")
        started_ok = True
    except Exception as e:
        wlog(f"Popen 启动失败: {e}")
        try:
            # 兜底：os.startfile
            if hasattr(os, 'startfile'):
                os.startfile(old_exe)
                wlog("已启动新程序 (os.startfile 兜底)")
                started_ok = True
        except Exception as e2:
            wlog(f"startfile 启动也失败: {e2}")

    # ===== 清理临时文件 =====
    for p in [new_exe, bak_exe] + cleanups:
        try:
            if os.path.exists(p):
                os.remove(p)
                wlog(f"清理: {os.path.basename(p)}")
        except Exception as e:
            wlog(f"清理 {os.path.basename(p)} 失败: {e}")

    wlog(f"更新完成，新程序启动状态: {'OK' if started_ok else '失败(请手动启动)'}")
    return 0


def gui_main():
    """GUI模式入口"""
    root = Tk()
    app = DailyReportGUI(root)
    root.mainloop()

def main():
    """
    主入口（按优先级判断模式）:
      1) --update-worker=<jsonPath>   纯后台覆盖模式（优先级最高，绝不加载GUI）
      2) -c / --cli                   命令行模式
      3) 默认                          GUI 模式
    """
    for arg in sys.argv[1:]:
        if arg.startswith('--update-worker='):
            _path = arg.split('=', 1)[1].strip('"')
            try:
                _rc = update_worker_main(_path)
            except Exception as _e:
                try:
                    import tempfile as _tf
                    _log = os.path.join(_tf.gettempdir(), 'update_last.log')
                    with open(_log, 'a', encoding='utf-8') as _f:
                        import traceback as _tb
                        _f.write(f"[{datetime.now():%H:%M:%S}] worker崩溃: {_e}\n{_tb.format_exc()}\n")
                except Exception:
                    pass
                _rc = 99
            sys.exit(_rc)
    if '-c' in sys.argv or '--cli' in sys.argv:
        cli_main()
    else:
        gui_main()

if __name__ == '__main__':
    main()