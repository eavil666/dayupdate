#!/usr/bin/env python3
"""
值守日报业务域（方案一拆分）
功能：告警文件加载/分类、威胁等级统计、docx 值守日报渲染（八大板块）、情报/跟进/总结。
依赖：common（日志/路径）、ipdb（is_private_ip/is_excluded_ip/load_config/load_terminal_ip_table）。
注意：classify 等 IP 判定逻辑在 ipdb（IP 判定为 IP 域职责）；本模块只做业务统计与渲染。
"""

import ipaddress
import os
import re
import tempfile
from datetime import datetime
from pathlib import Path

from common import _log, runtime_dir
from ipdb import is_excluded_ip, is_private_ip, load_config


def classify(ip, zone, geo, conf):
    try:
        addr = ipaddress.ip_address(str(ip).strip())
    except ValueError:
        return "未知"
    for net in conf["nets"]:
        if addr in net:
            return "内网"
    # 使用None检查替代pd.isna
    zone = "" if zone is None or (hasattr(zone, "__class__") and zone.__class__.__name__ == "NaNType") else str(zone)
    geo = "" if geo is None or (hasattr(geo, "__class__") and geo.__class__.__name__ == "NaNType") else str(geo)
    if zone in conf["zones"] and any(g in geo for g in conf["geos"]):
        return "内网"
    return "外网" if addr.is_global else "待确认"


def load_single_file(path, conf):
    import pandas as pd

    df = pd.read_excel(path, sheet_name=0, dtype={"源 IP": str, "目的 IP": str})
    df.columns = df.columns.str.strip()
    cols = []
    seen = set()
    for c in df.columns:
        if c in seen:
            i = 1
            while f"{c}_{i}" in seen:
                i += 1
            cols.append(f"{c}_{i}")
            seen.add(f"{c}_{i}")
        else:
            cols.append(c)
            seen.add(c)
    df.columns = cols
    col_map = {}
    target_counts = {}
    for c in df.columns:
        base = c.split("_")[0] if "_" in c else c
        target = None
        if "源" in base and "IP" in base and "目的" not in base:
            target = "源IP"
        elif "目的" in base and "IP" in base:
            target = "目的IP"
        elif "攻击" in base and ("名称" in base or "类型" in base):
            target = "攻击名称"
        elif "威胁" in base and "等级" in base:
            target = "威胁等级"
        elif "源" in base and "区域" in base:
            target = "源区域"
        elif "源" in base and "地理" in base:
            target = "源地理信息"
        elif "情报" in base or "IOC" in base.upper():
            target = "情报IOC"
        elif "攻击" in base and "阶段" in base:
            target = "攻击阶段"
        elif "攻击" in base and "状态" in base:
            target = "攻击状态"
        if target:
            if target in target_counts:
                target_counts[target] += 1
                col_map[c] = f"{target}_{target_counts[target]}"
            else:
                target_counts[target] = 1
                col_map[c] = target
    df = df.rename(columns=col_map)
    for must in ["源IP", "目的IP", "攻击名称", "威胁等级"]:
        if must not in df.columns:
            df[must] = ""
    df["网络类型"] = df.apply(lambda r: classify(r["源IP"], r.get("源区域", ""), r.get("源地理信息", ""), conf), axis=1)
    return df


def load_and_classify(paths, conf):
    dfs = []
    import pandas as pd

    for path in paths:
        _log(f"[+] 读取文件: {path.name}")
        df = load_single_file(path, conf)
        dfs.append(df)
    if not dfs:
        raise ValueError("没有读取到任何数据")
    df = pd.concat(dfs, ignore_index=True)
    _log(f"[+] 合并后共 {len(df)} 条记录")
    return df


LEVELS = ["严重", "高危", "中危", "低危"]


def analyze(df):
    total = len(df)
    internal = df[df["网络类型"] == "内网"]
    external = df[df["网络类型"] == "外网"]

    def by_level(sub):
        return {lv: int((sub["威胁等级"] == lv).sum()) for lv in LEVELS}

    external_to_internal = external.copy()
    # 空 external 时跳过过滤：pandas 空 df 用空 bool 过滤会丢失全部列（KeyError）
    if "目的IP" in external.columns and len(external) > 0:
        external_to_internal = external[external["目的IP"].apply(is_private_ip)]
    external_to_internal = external_to_internal[~external_to_internal["源IP"].apply(is_excluded_ip)]
    ban_count = int(external_to_internal["源IP"].nunique()) if len(external_to_internal) > 0 else 0
    return {
        "total": total,
        "internal": internal,
        "external": external,
        "int_count": len(internal),
        "ext_count": len(external),
        "int_level": by_level(internal),
        "ext_level": by_level(external),
        "ban_count": ban_count,
    }


def _set_run_font(run, size=None, bold=None):
    """统一设置 run 字体：英文 Times New Roman，中文宋体"""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _parse_lines(text, skip_example=True):
    """解析多行文本：按换行分割，去除空行和序号前缀"""
    if not text or not text.strip():
        return []
    items = []
    for line in text.split("\n"):
        line = line.strip()
        if line and (not skip_example or not line.startswith("示例：")):
            line = re.sub(r"^\d+\.\s*", "", line)
            items.append(line)
    return items


def _add_para(doc, text, bold=False, size=None):
    """添加段落并统一字体"""
    p = doc.add_paragraph()
    r = p.add_run(str(text))
    _set_run_font(r, size=size, bold=bold)
    return p


def _add_heading(doc, text, level=1):
    """添加标题并统一字体"""
    h = doc.add_heading(text, level)
    for run in h.runs:
        _set_run_font(run)
    return h


def _add_numbered_list(doc, items, start=1):
    """添加编号列表段落"""
    for i, item in enumerate(items, start):
        _add_para(doc, f"{i}. {item}")


def _add_table(doc, headers, widths, rows):
    """创建表格并填充数据，统一设置样式和列宽"""
    from docx.shared import Pt

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    _hdr(tbl, headers)
    for row_data in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            _set_cell(cells[i], val)
    _fit_table(tbl, [Pt(w) for w in widths])
    return tbl


def _hdr(table, headers):
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(p.add_run(h), size=9, bold=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_cell(cell, text, bold=False, size=9, padding=None):
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    _set_run_font(p.add_run(str(text)), size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 仅当指定 padding 时才设置，否则使用 Word 默认值
    if padding is not None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = tcPr.find(qn("w:tcMar"))
        if tcMar is None:
            tcMar = OxmlElement("w:tcMar")
            tcPr.append(tcMar)
        for child in list(tcMar):
            tcMar.remove(child)
        for side, val in padding.items():
            elem = OxmlElement(f"w:{side}")
            elem.set(qn("w:w"), val)
            elem.set(qn("w:type"), "dxa")
            tcMar.append(elem)


def _fit_table(table, widths=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if widths is None:
        widths = []
    # 在所有行上设置单元格宽度
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]
    # 设置表格固定布局（与用户调整的文档一致）
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")
    # 允许跨页断行
    for row in table.rows:
        trPr = row._tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            row._tr.insert(0, trPr)
        cant_split = trPr.find(qn("w:cantSplit"))
        if cant_split is None:
            cant_split = OxmlElement("w:cantSplit")
            trPr.append(cant_split)


def _lookup_geo(ip):
    """IP 归属查询（薄封装）：委托 ipdb.lookup_ip_geo 三级链路
    （geo_cache.json -> ip2region 离线 -> pconline 在线补全），IP 归属表与日报共用。
    """
    from ipdb import lookup_ip_geo

    return lookup_ip_geo(ip)


def _render_level_chart(stats, save_path):
    """用 PIL 手绘威胁等级柱状图（内网/外网对比），失败返回 None。

    不引入 matplotlib（exe 体积膨胀），PIL 画简单柱状图足够日报使用。
    中文字体优先微软雅黑/黑体/宋体，找不到则用默认字体（中文可能方块，但不崩）。
    """
    try:
        from PIL import Image, ImageDraw, ImageFont

        levels = list(LEVELS)  # 严重/高危/中危/低危
        int_vals = [stats["int_level"].get(lv, 0) for lv in levels]
        ext_vals = [stats["ext_level"].get(lv, 0) for lv in levels]

        W, H = 660, 360
        margin_l, margin_r, margin_t, margin_b = 70, 20, 50, 60
        plot_w, plot_h = W - margin_l - margin_r, H - margin_t - margin_b

        img = Image.new("RGB", (W, H), "white")
        d = ImageDraw.Draw(img)

        font_path = None
        for fp in (
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\simsun.ttc",
        ):
            if os.path.exists(fp):
                font_path = fp
                break
        try:
            f_title = ImageFont.truetype(font_path, 16) if font_path else ImageFont.load_default()
            f_label = ImageFont.truetype(font_path, 12) if font_path else ImageFont.load_default()
            f_axis = ImageFont.truetype(font_path, 10) if font_path else ImageFont.load_default()
        except Exception:
            f_title = f_label = f_axis = ImageFont.load_default()

        # 标题
        d.text((W // 2, 18), "威胁等级分布（内网 / 外网）", font=f_title, fill="black", anchor="mm")

        # 坐标轴
        d.line([(margin_l, margin_t), (margin_l, H - margin_b)], fill="black", width=1)
        d.line([(margin_l, H - margin_b), (W - margin_r, H - margin_b)], fill="black", width=1)

        max_val = max(max(int_vals), max(ext_vals), 1)
        n = len(levels)
        group_w = plot_w / n
        bar_w = min(22, group_w * 0.28)

        for i, lv in enumerate(levels):
            cx = margin_l + group_w * i + group_w / 2
            # 内网柱（蓝）
            iv = int_vals[i]
            ih = plot_h * iv / max_val
            d.rectangle(
                [cx - bar_w - 2, H - margin_b - ih, cx - 2, H - margin_b],
                fill="#4472C4",
            )
            if iv > 0:
                d.text((cx - bar_w / 2 - 2, H - margin_b - ih - 14), str(iv), font=f_axis, fill="#4472C4", anchor="mm")
            # 外网柱（红）
            ev = ext_vals[i]
            eh = plot_h * ev / max_val
            d.rectangle(
                [cx + 2, H - margin_b - eh, cx + bar_w + 2, H - margin_b],
                fill="#C00000",
            )
            if ev > 0:
                d.text((cx + bar_w / 2 + 2, H - margin_b - eh - 14), str(ev), font=f_axis, fill="#C00000", anchor="mm")
            # 类别标签
            d.text((cx, H - margin_b + 14), lv, font=f_label, fill="black", anchor="mm")

        # 图例
        d.rectangle([margin_l, 30, margin_l + 14, 44], fill="#4472C4")
        d.text((margin_l + 18, 37), "内网", font=f_axis, fill="black", anchor="lm")
        d.rectangle([margin_l + 60, 30, margin_l + 74, 44], fill="#C00000")
        d.text((margin_l + 78, 37), "外网", font=f_axis, fill="black", anchor="lm")

        img.save(save_path)
        return save_path
    except Exception:
        return None


def _render_attack_chart(attack_df, save_path, title=None, color="#C00000"):
    """PIL 手绘攻击类型 Top 横向条形图（名称长，横条更清晰），失败返回 None。

    attack_df：含"攻击名称"列的告警 DataFrame（外网或内网均可）
    title：图标题（默认"外网攻击类型分布 TOP8"）
    color：柱体颜色（外网红 / 内网蓝）
    """
    try:
        from PIL import Image, ImageDraw, ImageFont

        counts = attack_df["攻击名称"].value_counts().head(8)
        if len(counts) == 0:
            return None
        items = [(str(k), int(v)) for k, v in counts.items()]

        W, H = 660, 60 + len(items) * 34
        margin_l, margin_r, margin_t, margin_b = 220, 60, 40, 24
        plot_w = W - margin_l - margin_r
        bar_h = 20
        row_h = 34

        img = Image.new("RGB", (W, H), "white")
        d = ImageDraw.Draw(img)

        font_path = None
        for fp in (
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\simsun.ttc",
        ):
            if os.path.exists(fp):
                font_path = fp
                break
        try:
            f_title = ImageFont.truetype(font_path, 16) if font_path else ImageFont.load_default()
            f_label = ImageFont.truetype(font_path, 11) if font_path else ImageFont.load_default()
            f_axis = ImageFont.truetype(font_path, 10) if font_path else ImageFont.load_default()
        except Exception:
            f_title = f_label = f_axis = ImageFont.load_default()

        d.text((W // 2, 18), title or "外网攻击类型分布 TOP8", font=f_title, fill="black", anchor="mm")

        max_val = max(v for _, v in items)
        # 坐标轴
        d.line([(margin_l, margin_t), (margin_l, H - margin_b)], fill="black", width=1)
        d.line([(margin_l, H - margin_b), (W - margin_r, H - margin_b)], fill="black", width=1)

        for i, (name, v) in enumerate(items):
            y = margin_t + i * row_h
            bw = plot_w * v / max_val
            d.rectangle([margin_l, y, margin_l + bw, y + bar_h], fill=color)
            if v > 0:
                d.text((margin_l + bw + 6, y + bar_h / 2), str(v), font=f_axis, fill=color, anchor="lm")
            # 名称（右侧截断，避免超宽）
            label = name if len(name) <= 16 else name[:15] + "…"
            d.text((margin_l - 8, y + bar_h / 2), label, font=f_label, fill="black", anchor="rm")

        img.save(save_path)
        return save_path
    except Exception:
        return None


def render(
    conf,
    df,
    stats,
    health_rows,
    intel_list,
    date,
    out_path,
    work_summary=None,
    follow_items=None,
    intel_items=None,
    threat_data=None,
):
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = docx.Document()

    # 统一设置样式字体
    def _apply_style_font(style):
        f = style.font
        f.name = "Times New Roman"
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(attr), "Times New Roman")
        rFonts.set(qn("w:eastAsia"), "宋体")

    _apply_style_font(doc.styles["Normal"])
    doc.styles["Normal"].font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2", "Title"):
        try:
            _apply_style_font(doc.styles[name])
        except KeyError:
            pass
    h = _add_heading(doc, conf["title"], 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    r = sub.add_run(f"日期：{date[:4]}-{date[4:6]}-{date[6:]}    编制：网络安全值守组    密级：内部")
    _set_run_font(r, bold=True)
    sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 一、当日态势概览与重点工作总结
    _add_heading(doc, "一、当日态势概览与重点工作总结", 1)
    total = max(stats["total"], 1)
    auto_summary = (
        f"今日共捕获告警 {stats['total']} 起，其中内网 {stats['int_count']} 起"
        f"（{stats['int_count'] / total * 100:.1f}%），外网 {stats['ext_count']} 起"
        f"（{stats['ext_count'] / total * 100:.1f}%），累计处置 IP {stats['ban_count']} 个。"
    )
    _add_para(doc, f"1. {auto_summary}")
    work_items = _parse_lines(work_summary)
    _add_numbered_list(doc, work_items, start=2)

    # 二、关键指标
    _add_heading(doc, "二、关键指标", 1)
    _add_table(
        doc,
        ["内网告警", "次数", "外网告警", "次数"],
        [110, 111, 110, 111],
        [(lv, stats["int_level"].get(lv, 0), lv, stats["ext_level"].get(lv, 0)) for lv in LEVELS],
    )
    # 威胁等级柱状图（PIL 手绘，失败静默跳过）
    _chart_path = _render_level_chart(stats, os.path.join(tempfile.gettempdir(), f"report_chart_{date}.png"))
    if _chart_path:
        try:
            doc.add_picture(_chart_path, width=docx.shared.Cm(15))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
    for idx, (title, body) in enumerate(
        [
            ("可疑域名解析请求分析", "通过定期对应用进行安全事件监测，暂未发现可疑域名解析请求事件。"),
            ("敏感文件分析", "通过定期对应用进行安全事件监测，暂未发现敏感文件。"),
            ("黑链挂马分析", "通过定期对应用进行安全事件监测，暂未发现黑链/挂马事件。"),
            ("可用性分析", "通过定期对应用进行安全事件监测，暂未发现存在可用性事件。"),
            ("篡改分析", "通过定期对应用进行安全事件监测，暂未发现网页篡改事件。"),
            ("域名劫持分析", "通过定期对应用进行安全事件监测,暂未发现域名劫持事件。"),
        ],
        start=1,
    ):
        _add_heading(doc, f"{idx}. {title}", 2)
        _add_para(doc, body)

    # 三、研判与处置流程
    _add_heading(doc, "三、研判与处置流程", 1)
    for s in [
        "流量采集：安全设备日志导出 + NDR 探针 + 防火墙 IPS",
        "情报比对：威胁情报平台 IOC 命中 / ASN 与归属地归类",
        "内外网判定：源 IP 静态段 + 资产归属 + 流量方向三重校验",
        "分级处置：P0 即时处置 / P1 当日处置 / P2 观察 / P3 沉淀",
    ]:
        p = doc.add_paragraph(style="List Number")
        _set_run_font(p.add_run(s))

    # 四、资产健康检查
    _add_heading(doc, "四、资产健康检查", 1)
    _add_para(doc, "值守期间持续监控安全设备运行状态，确保监测能力不降级。")
    _add_table(
        doc,
        ["探针名称", "IP", "状态", "备注"],
        [120, 100, 111, 111],
        [(name, ip, "正常", "") for name, ip, *_ in health_rows],
    )
    _add_para(doc, f"日志留存：{conf['retention']} 天。")

    # 五、外网攻击研判与处置
    _add_heading(doc, "五、外网攻击研判与处置", 1)
    ext = stats["external"]
    bad_ips, threat_sources = threat_data if threat_data else (set(), {})
    threat_hits = []  # 命中公开威胁名单的源 IP（用于处置建议）
    if len(ext) > 0:
        grp = ext.groupby("攻击名称", sort=False).agg({"威胁等级": "first", "源IP": "count"}).reset_index()
        ext_rows = []
        for idx, (_, row) in enumerate(grp.iterrows(), start=1):
            # 该攻击类型下的源 IP 集合
            type_ips = ext.loc[ext["攻击名称"] == row["攻击名称"], "源IP"].dropna().astype(str).tolist()
            hit_ips = [ip for ip in type_ips if ip in bad_ips]
            threat_txt = f"High×{len(hit_ips)}" if hit_ips else "—"
            ext_rows.append((idx, row["攻击名称"], f"{int(row['源IP'])} 起", threat_txt, "流量特征+情报比对", "已处置"))
        # 全部外网源 IP 中命中威胁名单的（去重）
        all_ext_ips = ext["源IP"].dropna().astype(str).unique().tolist()
        for ip in all_ext_ips:
            if ip in bad_ips:
                hits = threat_sources.get(ip, [])
                threat_hits.append((ip, ",".join(hits)))
    else:
        ext_rows = [("（无外网告警）", "", "", "", "", "")]
    _add_table(
        doc,
        ["序号", "威胁类型", "命中次数", "威胁源分级", "研判依据", "状态"],
        [41, 173, 51, 71, 106, 41],
        ext_rows,
    )
    # 攻击类型分布图（PIL 横向条形图，失败静默跳过）
    _attack_chart = _render_attack_chart(ext, os.path.join(tempfile.gettempdir(), f"attack_chart_{date}.png"))
    if _attack_chart:
        try:
            doc.add_picture(_attack_chart, width=docx.shared.Cm(15))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    # 六、内网异常研判与处置
    _add_heading(doc, "六、内网异常研判与处置", 1)
    intdf = stats["internal"]
    if len(intdf) > 0:
        agg_map = {"源IP": "count"}
        if "威胁等级" in intdf.columns:
            agg_map["威胁等级"] = "first"
        grp = intdf.groupby("攻击名称", sort=False).agg(agg_map).reset_index()
        int_rows = []
        for idx, (_, row) in enumerate(grp.iterrows(), start=1):
            # 该攻击类型下的源 IP 集合（威胁名单匹配，口径与内网一致）
            type_ips = intdf.loc[intdf["攻击名称"] == row["攻击名称"], "源IP"].dropna().astype(str).tolist()
            hit_ips = [ip for ip in type_ips if ip in bad_ips]
            threat_txt = f"High×{len(hit_ips)}" if hit_ips else "—"
            int_rows.append((idx, row["攻击名称"], f"{int(row['源IP'])} 起", threat_txt, "流量特征+情报比对", "已处置"))
    else:
        int_rows = [("（无内网告警）", "", "", "", "", "")]
    _add_table(
        doc,
        ["序号", "威胁类型", "命中次数", "威胁源分级", "研判依据", "状态"],
        [41, 173, 51, 71, 106, 41],
        int_rows,
    )
    # 内网攻击类型分布图（PIL 横向条形图，蓝色区分外网；失败静默跳过）
    _int_chart = _render_attack_chart(
        intdf,
        os.path.join(tempfile.gettempdir(), f"int_chart_{date}.png"),
        title="内网攻击类型分布 TOP8",
        color="#4472C4",
    )
    if _int_chart:
        try:
            doc.add_picture(_int_chart, width=docx.shared.Cm(15))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
    # 内网外联威胁研判：内网主机主动外联（源内网->目的外部），
    # 目的命中公开威胁名单 = 疑似木马回连/挖矿（失陷信号）
    _add_heading(doc, "内网外联威胁研判", 2)
    _ext_dst = intdf.copy()
    if len(_ext_dst) > 0 and "目的IP" in _ext_dst.columns:

        def _is_ext_ip(ip):
            try:
                return not ipaddress.ip_address(str(ip).strip()).is_private
            except ValueError:
                return False

        _ext_dst = _ext_dst[_ext_dst["目的IP"].notna() & _ext_dst["目的IP"].apply(_is_ext_ip)].copy()
        # 外联目的聚合（源主机 x 目的IP）
        _geo_col = "目的地理信息" if "目的地理信息" in _ext_dst.columns else None
        _grp = (
            _ext_dst.groupby(["源IP", "目的IP"])
            .agg(
                count=("目的IP", "size"),
                geo=(_geo_col, "first") if _geo_col else ("目的IP", lambda x: ""),
            )
            .reset_index()
        )
        # 常见公共服务（DNS等），外联属正常：先排除再取 Top，聚焦可疑外联
        _public_svc = {"114.114.114.114", "223.5.5.5", "8.8.8.8", "1.1.1.1", "119.29.29.29", "180.76.76.76"}
        _grp = _grp[~_grp["目的IP"].astype(str).str.strip().isin(_public_svc)]
        _grp = _grp.sort_values("count", ascending=False).head(15)
        _ext_rows = []
        for _idx, (_, _row) in enumerate(_grp.iterrows(), start=1):
            _dst = str(_row["目的IP"]).strip()
            _geo = _lookup_geo(_dst) or (str(_row["geo"]) if _row["geo"] is not None else "")
            _geo = "" if _geo.lower() in ("nan", "none") else _geo[:20]
            if _dst in bad_ips:
                _lvl = "High"
                _judge = "疑回连/挖矿，建议隔离"
            elif int(_row["count"]) >= 3:
                _lvl = "观察"
                _judge = "外联频繁，核查业务"
            else:
                _lvl = "观察"
                _judge = "外联行为，业务核查"
            _ext_rows.append((_idx, _row["源IP"], _row["目的IP"], _geo, int(_row["count"]), _lvl, _judge))
        # 列宽总 450pt（页面约 15cm 可放下）；研判列 130pt 装下 11 字中文，Geo 大多空给窄
        _add_table(
            doc,
            ["序号", "内网主机", "外联目的IP", "目的归属", "次数", "威胁分级", "研判"],
            [25, 80, 95, 45, 30, 45, 130],
            _ext_rows,
        )
    else:
        _add_para(doc, "今日无内网外联行为。")

    # 七、重点事件研判
    _add_heading(doc, "七、重点事件研判", 1)
    key = df[df["威胁等级"].isin(conf["crit_levels"])].copy()
    if len(key) > 0:
        key["_p"] = key["威胁等级"].map({lv: i for i, lv in enumerate(LEVELS)})
        key = key.sort_values("_p")
        grp = key.groupby(["攻击名称", "源IP"], sort=False).size().reset_index(name="count")
        grp = grp.sort_values("count", ascending=False).head(conf["top"])
        evt_name_map = {}
        for idx, row in enumerate(grp[["攻击名称"]].drop_duplicates().values, 1):
            evt_name_map[row[0]] = idx
        key_rows = [
            (evt_name_map[row["攻击名称"]], row["攻击名称"], row["源IP"], row["count"]) for _, row in grp.iterrows()
        ]
        _add_table(doc, ["序号", "事件名称", "源地址", "攻击次数"], [41, 145, 135, 121], key_rows)
    else:
        _add_para(doc, "今日无严重/高危级事件。")
    # 自动研判结论（结合威胁分级）
    if len(ext) > 0:
        top_att = ext["攻击名称"].value_counts().head(1)
        att_name = top_att.index[0] if len(top_att) else ""
        att_n = int(top_att.iloc[0]) if len(top_att) else 0
        hit_n = len(threat_hits)
        if hit_n:
            conclusion = (
                f"研判结论：今日外网攻击以「{att_name}」为主（{att_n} 起），"
                f"其中 {hit_n} 个源 IP 命中公开威胁名单（多为境外 VPS 扫描源），"
                f"建议对上述源实施临时封禁并持续监控；内网侧未发现失陷迹象。"
            )
        else:
            conclusion = (
                f"研判结论：今日外网攻击以「{att_name}」为主（{att_n} 起），"
                f"源 IP 均未命中公开威胁名单，判定为常规扫描探测，已按流程处置，持续观察。"
            )
        _add_para(doc, conclusion)

    # 八、情报动态
    _add_heading(doc, "八、情报动态", 1)
    _add_para(doc, "当日需关注的新增 CVE / 行业预警：")
    intel_parsed = _parse_lines(intel_items, skip_example=False)
    if intel_parsed:
        _add_numbered_list(doc, intel_parsed)
    elif intel_list:
        _add_table(
            doc,
            ["类型", "编号", "风险", "关联资产", "应对/时限"],
            [60, 85, 65, 115, 117],
            [
                (
                    item.get("类型", ""),
                    item.get("编号", ""),
                    item.get("风险", ""),
                    item.get("关联资产", ""),
                    f"{item.get('应对', '')} / {item.get('时限', '')}",
                )
                for item in intel_list
            ],
        )
    else:
        _add_para(doc, "（暂无新增情报，详见威胁情报平台）")

    # 九、待跟进事项
    _add_heading(doc, "九、待跟进事项", 1)
    default_items = [
        "内网高危告警溯源与处置闭环",
        "外网封禁 IP 清单同步至边界防火墙",
        "失陷终端取证与隔离",
        "资产健康检查异常处理",
    ]
    follow_parsed = _parse_lines(follow_items, skip_example=False)
    _add_numbered_list(doc, follow_parsed if follow_parsed else default_items)

    doc.save(out_path)


def load_intel(conf):
    p = Path(os.path.join(runtime_dir, conf["intel_file"]))
    if not p.exists():
        return []
    try:
        import pandas as pd

        return pd.read_csv(p).fillna("").to_dict("records")
    except Exception as e:
        _log(f"[!] 加载情报文件失败: {e}")
        return []


def pick_input_and_date(pattern):
    # 文件名排除关键词（不区分大小写）：非安全告警类的辅助/输出/配置文件
    _EXCLUDE_KEYS = (
        "终端ip地址表",
        "终端ip",
        "业务ip",
        "业务IP",
        "biz_ip",
        "ip归属分析",
        "IP归属分析",
        "ip分析",
        "IP分析",
        "值守保障日报",
        "值守日报",
        "日报",
        "日报汇总",
        "config",
        "配置",
        "config.ini",
        "version",
        "情报",
        "intel",
    )

    def _is_excluded(fpath):
        s = fpath.stem
        for k in _EXCLUDE_KEYS:
            if k.lower() in s.lower():
                return k
        return None

    files = sorted(Path(runtime_dir).glob(pattern), key=os.path.getmtime, reverse=True)
    if not files:
        files = sorted(Path(runtime_dir).glob("*.xls*"), key=os.path.getmtime, reverse=True)

    # 先输出候选与排除情况，便于定位
    excluded_info = []
    filtered = []
    for f in files:
        reason = _is_excluded(f)
        if reason:
            excluded_info.append(f"  ! {f.name}  已排除(命中关键词: {reason})")
        else:
            filtered.append(f)
    files = filtered
    if excluded_info:
        _log("[*] 自动识别排除文件:")
        _log("\n".join(excluded_info))

    if not files:
        raise FileNotFoundError(
            "未找到安全告警类 Excel，请把安全设备导出的日志放到脚本目录，"
            "并确保文件名不含「业务IP/终端IP/值守日报/IP归属分析」等排除关键词"
        )

    # 锚点优先选择: stem 中包含 8 位日期 的最新 mtime 文件（真实告警导出一般带时间戳）
    dated_files = [f for f in files if re.search(r"(\d{8})", f.stem)]
    if dated_files:
        anchor = sorted(dated_files, key=os.path.getmtime, reverse=True)[0]
    else:
        anchor = files[0]  # 退化: 所有文件都没带日期，取 mtime 最新

    stem = anchor.stem
    m = re.search(r"(\d{8})", stem)
    if m:
        date = m.group(1)
        _log(f"[*] 自动识别锚点文件: {anchor.name} (日期: {date})")
    else:
        date = datetime.now().strftime("%Y%m%d")
        _log(f"[*] 自动识别锚点文件: {anchor.name} (文件名无日期，使用今日: {date})")

    target_files = []
    for f in files:
        fm = re.search(r"(\d{8})", f.stem)
        if fm and fm.group(1) == date:
            target_files.append(f)
    if not target_files:
        target_files = [anchor]
    _log(f"[*] 自动识别结果: 共 {len(target_files)} 个文件 (日期: {date})")
    return target_files, date


def generate_daily_report(files, date, work_summary=None, follow_items=None, intel_items=None, local_geos=None):
    conf = load_config(files, local_geos=local_geos)  # files：自动提取区域/归属地；local_geos：GUI 自定义
    df = load_and_classify(files, conf)
    stats = analyze(df)
    _log(
        f"[+] 总告警 {stats['total']} | 内网 {stats['int_count']} | 外网 {stats['ext_count']} | 处置 {stats['ban_count']}"
    )
    # 威胁名单（threat_check：磁盘缓存 6h，失败降级空集，不阻塞日报）
    threat_data = (set(), {})
    try:
        from threat_check import load_bad_ips

        _cache_file = Path(runtime_dir) / "threat_feeds_cache.json"
        bad_ips, sources = load_bad_ips(str(_cache_file))
        if bad_ips:
            threat_data = (bad_ips, sources)
            _log(f"[+] 威胁名单加载: {len(bad_ips)} 条（用于外网研判分级）")
        else:
            _log("[!] 威胁名单为空，外网研判表威胁源分级显示 —")
    except Exception as _e:
        _log(f"[!] 威胁名单加载失败: {_e}")
    health_rows = conf["probes"]
    intel_list = load_intel(conf)
    # 使用 runtime_dir（脚本/exe所在目录）作为输出基础目录
    out_dir = Path(runtime_dir) / conf["out_dir"]
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / f"值守保障日报{date}.docx"
    render(
        conf,
        df,
        stats,
        health_rows,
        intel_list,
        date,
        str(out_path),
        work_summary,
        follow_items,
        intel_items,
        threat_data,
    )
    _log(f"[✓] 值守日报已生成: {out_path}")
    return out_path
